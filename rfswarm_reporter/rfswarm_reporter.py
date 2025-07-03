#!/usr/bin/python
#
# 	Robot Framework Swarm
# 		Reporter
#    Version 1.6.0
#

import argparse
import base64  # used for embedding images  # used for xhtml export
import configparser
import difflib
import importlib.metadata
import json
import os
import platform
import shutil
import signal
import sqlite3
import sys
import threading
import time
import tkinter as tk  # python3
import zoneinfo  # says Requires python 3.9
from copy import copy  # used for xlsx export
from datetime import datetime  # , timezone
from io import BytesIO  # used for embedding images  # used for xhtml export
from typing import Any
import matplotlib  # required for matplot graphs
import matplotlib.font_manager as font_manager
import openpyxl  # used for xlsx export
import yaml
from docx import Document  # used for docx export
from docx.enum.style import WD_STYLE_TYPE  # used for docx export
from docx.enum.text import WD_ALIGN_PARAGRAPH  # used for docx export
from docx.oxml.shared import OxmlElement, qn  # used for docx export
from docx.shared import Cm, Pt, RGBColor  # used for docx export
from lxml import etree  # used for xhtml export
from lxml.builder import E, ElementMaker  # used for xhtml export

# required for matplot graphs
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas

# required for matplot graphs
from matplotlib.figure import Figure  # required for matplot graphs
from PIL import Image

if True:  # noqa: E402
	sys.path.append(os.path.abspath(os.path.dirname(__file__)))
	from ReporterBase import ReporterBase
	from ReporterGUI import ReporterGUI


matplotlib.use("TkAgg") 	# required for matplot graphs

__name__ = "rfswarm-manager"


class ReporterCore:

	cg_data: Any = {}
	t_export: Any = {}

	def __init__(self, master=None):
		base.debugmsg(0, "Robot Framework Swarm: Reporter")
		base.debugmsg(0, "	Version", base.version)
		signal.signal(signal.SIGINT, self.on_closing)

		font_manager._get_fontconfig_fonts.cache_clear()

		base.debugmsg(9, "ArgumentParser")
		# Check for command line args
		parser = argparse.ArgumentParser()
		parser.add_argument('-g', '--debug', help='Set debug level, default level is 0')
		parser.add_argument('-v', '--version', help='Display the version and exit', action='store_true')
		parser.add_argument('-i', '--ini', help='path to alternate ini file')
		parser.add_argument('-n', '--nogui', help='Don\'t display the GUI', action='store_true')
		parser.add_argument('-d', '--dir', help='Results directory')
		parser.add_argument('-t', '--template', help='Specify the template')
		parser.add_argument('--html', help='Generate a HTML report', action='store_true')
		# parser.add_argument('--pdf', help='Generate a PDF report', action='store_true')
		parser.add_argument('--docx', help='Generate a MS Word report', action='store_true')
		parser.add_argument('--xlsx', help='Generate a MS Excel report', action='store_true')
		# parser.add_argument('--odt', help='Generate an OpenOffice/LibreOffice Writer report', action='store_true')
		# parser.add_argument('--ods', help='Generate an OpenOffice/LibreOffice Calc report', action='store_true')
		parser.add_argument('-c', '--create', help='ICON : Create application icon / shortcut')
		base.args = parser.parse_args()
		base.core = self

		base.debugmsg(6, "base.args: ", base.args)

		if base.args.debug:
			base.debuglvl = int(base.args.debug)

		if base.args.version:
			self.show_additional_versions()
			exit()

		if base.args.create:
			if base.args.create.upper() in ["ICON", "ICONS"]:
				self.create_icons()
			else:
				base.debugmsg(0, "create with option ", base.args.create.upper(), "not supported.")
			exit()

		if base.args.nogui:
			base.displaygui = False

		base.debugmsg(6, "ConfigParser")
		base.config = configparser.ConfigParser()

		#
		# 	ensure ini file
		#
		base.reporter_ini = base.findiniloctaion()

		if base.args.ini:
			base.save_ini = False
			base.debugmsg(5, "base.args.ini: ", base.args.ini)
			base.reporter_ini = base.args.ini

		if os.path.isfile(base.reporter_ini):
			base.debugmsg(7, "reporter_ini: ", base.reporter_ini)
			arrconfigfile = os.path.splitext(base.reporter_ini)
			base.debugmsg(5, "arrconfigfile: ", arrconfigfile)
			if len(arrconfigfile) < 2:
				base.debugmsg(0, "Configuration file ", base.reporter_ini, " missing extention, unable to determine supported format. Plesae use extentions .ini, .yaml or .json")
				exit()
			if arrconfigfile[1].lower() not in [".ini", ".yml", ".yaml", ".json"]:
				base.debugmsg(0, "Configuration file ", base.reporter_ini, " has an invalid extention, unable to determine supported format. Plesae use extentions .ini, .yaml or .json")
				exit()
			if arrconfigfile[1].lower() == ".ini":
				base.config.read(base.reporter_ini, encoding="utf8")
			else:
				configdict = {}
				if arrconfigfile[1].lower() in [".yml", ".yaml"]:
					# read yaml file
					base.debugmsg(5, "read yaml file")
					with open(base.reporter_ini, 'r', encoding="utf-8") as f:
						configdict = yaml.safe_load(f)
						configdict = base.configparser_safe_dict(configdict)
						base.debugmsg(5, "configdict: ", configdict)
				if arrconfigfile[1].lower() == ".json":
					# read json file
					base.debugmsg(5, "read json file")
					with open(base.reporter_ini, 'r', encoding="utf-8") as f:
						configdict = json.load(f)
						configdict = base.configparser_safe_dict(configdict)
						base.debugmsg(5, "configdict: ", configdict)
				base.debugmsg(5, "configdict: ", configdict)
				base.config.read_dict(configdict)
		else:
			base.saveini()

		base.debugmsg(0, "	Configuration File: ", base.reporter_ini)

		base.debugmsg(9, "base.config: ", base.config._sections)

		#
		# GUI
		#
		if 'GUI' not in base.config:
			base.config['GUI'] = {}
			base.saveini()

		if 'win_width' not in base.config['GUI']:
			base.config['GUI']['win_width'] = "800"
			base.saveini()

		if 'win_height' not in base.config['GUI']:
			base.config['GUI']['win_height'] = "390"
			base.saveini()

		if 'donation_reminder' not in base.config['GUI']:
			base.config['GUI']['donation_reminder'] = "0"
			base.saveini()

		#
		# Reporter
		#

		if 'Reporter' not in base.config:
			base.config['Reporter'] = {}
			base.saveini()

		if 'ResultDir' not in base.config['Reporter']:
			base.config['Reporter']['ResultDir'] = base.dir_path
			base.saveini()
		else:
			if not os.path.isdir(base.config['Reporter']['ResultDir']):
				base.config['Reporter']['ResultDir'] = base.dir_path
				base.saveini()

		if 'Results' not in base.config['Reporter']:
			base.config['Reporter']['Results'] = ""
			base.saveini()
		else:
			if not os.path.isfile(base.config['Reporter']['Results']):
				base.config['Reporter']['Results'] = ""
				base.saveini()

		if 'Report' not in base.config['Reporter']:
			base.config['Reporter']['Report'] = ""
			base.saveini()
		else:
			if not os.path.isfile(base.config['Reporter']['Report']):
				base.config['Reporter']['Report'] = ""
				base.saveini()

		if 'Template' not in base.config['Reporter']:
			base.config['Reporter']['Template'] = ""
			base.saveini()
		else:
			if not os.path.isfile(base.config['Reporter']['Template']):
				base.config['Reporter']['Template'] = ""
				base.saveini()

		if 'TemplateDir' not in base.config['Reporter']:
			base.config['Reporter']['TemplateDir'] = ""
			base.saveini()
		else:
			if not os.path.isdir(base.config['Reporter']['TemplateDir']):
				base.config['Reporter']['TemplateDir'] = ""
				base.saveini()

		usetemplate = False
		if base.args.template:
			usetemplate = True
			base.config['Reporter']['Template'] = base.whitespace_set_ini_value(base.args.template)

		if base.args.dir:
			# do some sanity checks before blindly setting
			rdir = base.args.dir
			base.debugmsg(5, "rdir:", rdir)
			if os.path.exists(rdir):
				if os.path.isfile(rdir):
					rdir = os.path.dirname(rdir)
					base.debugmsg(5, "rdir:", rdir)
				dname = os.path.basename(rdir)
				dbfile = "{}.db".format(dname)
				dbpath = os.path.join(rdir, dbfile)
				base.debugmsg(5, "dbpath:", dbpath)
				if os.path.isfile(dbpath):
					base.config['Reporter']['Results'] = dbpath

		self.selectResults(base.config['Reporter']['Results'])

		if not usetemplate and "Report" in base.config['Reporter'] \
			and len(base.config['Reporter']['Report']) \
			and os.path.isfile(base.config['Reporter']['Report']):
			base.report_open()
		else:
			base.template_open(base.whitespace_get_ini_value(base.config['Reporter']['Template']))
			base.report_save()

		if base.args.html:
			# self.export_xhtml()
			self.t_export["xhtml"] = threading.Thread(target=self.export_xhtml)
			self.t_export["xhtml"].start()

		if base.args.docx:
			# self.export_word()
			self.t_export["docx"] = threading.Thread(target=self.export_word)
			self.t_export["docx"].start()

		if base.args.xlsx:
			# self.export_excel()
			self.t_export["xlsx"] = threading.Thread(target=self.export_excel)
			self.t_export["xlsx"].start()

		if base.displaygui:
			base.gui = ReporterGUI(base)
		else:
			# t_export
			for thd in self.t_export.keys():
				self.t_export[thd].join()
			self.on_closing()

	def show_additional_versions(self):

		base.debugmsg(0, "	Dependancy Versions")
		try:
			base.debugmsg(0, "		Python Version", sys.version)
		except Exception:
			pass

		try:
			base.debugmsg(0, "		SQLite Version", sqlite3.sqlite_version)
		except Exception:
			pass

		try:
			base.debugmsg(0, "		Tcl/Tk Version", tk.Tcl().call("info", "patchlevel"))
		except Exception:
			pass

	def mainloop(self):

		base.debugmsg(5, "mainloop start")

		if base.displaygui:
			base.gui.mainloop()

	def on_closing(self, _event=None, *extras):
		base.running = False
		base.debugmsg(5, "base.running:", base.running)

		base.debugmsg(5, "Close results db")
		# base.close_results_db()
		base.stop_db()

		time.sleep(1)
		base.debugmsg(2, "Exit")
		try:
			sys.exit(0)
		except SystemExit as e:
			try:
				remaining_threads = [t for t in threading.enumerate() if t is not threading.main_thread() and t.is_alive()]
				if remaining_threads:
					base.debugmsg(5, "Failed to gracefully exit RFSwarm-Reporter. Forcing immediate exit.")
					for thread in remaining_threads:
						base.debugmsg(9, "Thread name:", thread.name)
					os._exit(0)
				else:
					raise e

			except Exception as e:
				base.debugmsg(3, "Failed to exit with error:", e)
				os._exit(1)
		sys.stdout.flush()
		sys.stderr.flush()

	def create_icons(self):
		base.debugmsg(0, "Creating application icons for RFSwarm Reporter")
		appname = "RFSwarm Reporter"
		namelst = appname.split()
		base.debugmsg(6, "namelst:", namelst)
		projname = "-".join(namelst).lower()
		base.debugmsg(6, "projname:", projname)
		pipdata = importlib.metadata.distribution(projname)
		# print("files:", pipdata.files)
		# print("file0:", pipdata.files[0])
		reporter_executable = os.path.abspath(str(pipdata.locate_file(pipdata.files[0])))
		base.debugmsg(5, "reporter_executable:", reporter_executable)

		script_dir = os.path.dirname(os.path.abspath(__file__))
		base.debugmsg(5, "script_dir:", script_dir)
		icon_dir = os.path.join(pipdata.locate_file('rfswarm_reporter'), "icons")
		base.debugmsg(5, "icon_dir:", icon_dir)

		if platform.system() == 'Linux':
			fileprefix = "~/.local/share"
			if os.access("/usr/share", os.W_OK):
				try:
					base.ensuredir("/usr/share/applications")
					directoryfilename = os.path.join("/usr/share/applications", "rfswarm.directory")
					directorydata = ["test"]
					with open(directoryfilename, 'w') as df:
						df.writelines(directorydata)
					os.remove(directoryfilename)
					fileprefix = "/usr/share"
				except Exception:
					pass

			fileprefix = os.path.expanduser(fileprefix)

			# base.debugmsg(5, "Create .directory file")
			# directorydata = []
			# directorydata.append('[Desktop Entry]\n')
			# directorydata.append('Type=Directory\n')
			# directorydata.append('Name=RFSwarm\n')
			# directorydata.append('Icon=rfswarm-logo\n')
			#
			# directoryfilename = os.path.join(fileprefix, "desktop-directories", "rfswarm.directory")
			# directorydir = os.path.dirname(directoryfilename)
			# base.ensuredir(directorydir)
			#
			# base.debugmsg(5, "directoryfilename:", directoryfilename)
			# with open(directoryfilename, 'w') as df:
			# 	df.writelines(directorydata)
			#
			# directoryfilename = os.path.join(fileprefix, "applications", "rfswarm.directory")
			# directorydir = os.path.dirname(directoryfilename)
			# base.ensuredir(directorydir)
			# base.debugmsg(5, "directoryfilename:", directoryfilename)
			# with open(directoryfilename, 'w') as df:
			# 	df.writelines(directorydata)

			base.debugmsg(5, "Create .desktop file")
			desktopdata = []
			desktopdata.append('[Desktop Entry]\n')
			desktopdata.append('Name=' + appname + '\n')
			desktopdata.append('Exec=' + reporter_executable + '\n')
			desktopdata.append('Terminal=false\n')
			desktopdata.append('Type=Application\n')
			desktopdata.append('Icon=' + projname + '\n')
			desktopdata.append('Categories=RFSwarm;Development;\n')
			desktopdata.append('Keywords=rfswarm;reporter;\n')
			# desktopdata.append('\n')

			dektopfilename = os.path.join(fileprefix, "applications", projname + ".desktop")
			dektopdir = os.path.dirname(dektopfilename)
			base.ensuredir(dektopdir)

			base.debugmsg(5, "dektopfilename:", dektopfilename)
			with open(dektopfilename, 'w') as df:
				df.writelines(desktopdata)

			base.debugmsg(5, "Copy icons")
			# /usr/share/icons/hicolor/128x128/apps/
			# 	1024x1024  128x128  16x16  192x192  22x22  24x24  256x256  32x32  36x36  42x42  48x48  512x512  64x64  72x72  8x8  96x96
			# or
			#  ~/.local/share/icons/hicolor/256x256/apps/
			src_iconx128 = os.path.join(icon_dir, projname + "-128.png")
			base.debugmsg(5, "src_iconx128:", src_iconx128)
			dst_iconx128 = os.path.join(fileprefix, "icons", "hicolor", "128x128", "apps", projname + ".png")
			dst_icondir = os.path.dirname(dst_iconx128)
			base.ensuredir(dst_icondir)
			base.debugmsg(5, "dst_iconx128:", dst_iconx128)
			shutil.copy(src_iconx128, dst_iconx128)

			src_iconx128 = os.path.join(icon_dir, "rfswarm-logo-128.png")
			base.debugmsg(5, "src_iconx128:", src_iconx128)
			dst_iconx128 = os.path.join(fileprefix, "icons", "hicolor", "128x128", "apps", "rfswarm-logo.png")
			base.debugmsg(5, "dst_iconx128:", dst_iconx128)
			shutil.copy(src_iconx128, dst_iconx128)

		if platform.system() == 'Darwin':
			base.debugmsg(5, "Create folder structure in /Applications")

			src_iconx1024 = os.path.join(icon_dir, projname + "-1024.png")

			self.create_macos_app_bundle(appname, pipdata.version, reporter_executable, src_iconx1024)

		if platform.system() == 'Windows':
			base.debugmsg(5, "Create Startmenu shorcuts")
			roam_appdata = os.environ["APPDATA"]
			scutpath = os.path.join(roam_appdata, "Microsoft", "Windows", "Start Menu", appname + ".lnk")
			src_iconx128 = os.path.join(icon_dir, projname + "-128.ico")

			self.create_windows_shortcut(scutpath, reporter_executable, src_iconx128, "Performance testing with robot test cases", True)

	def create_windows_shortcut(self, scutpath, targetpath, iconpath, desc, minimised=False):
		pslst = []

		directorydir = os.path.dirname(scutpath)
		base.ensuredir(directorydir)

		pslst.append("$wshshell = New-Object -COMObject wscript.shell")
		pslst.append('$scut = $wshshell.CreateShortcut("""' + scutpath + '""")')
		pslst.append('$scut.TargetPath = """' + targetpath + '"""')
		pslst.append('$scut.IconLocation = """' + iconpath + '"""')
		if minimised:
			pslst.append("$scut.WindowStyle = 7")
		pslst.append("$scut.Description = '" + desc + "'")
		pslst.append("$scut.Save()")

		# psscript = '\n'.join(pslst)
		psscript = '; '.join(pslst)
		base.debugmsg(6, "psscript:", psscript)

		response = os.popen('powershell.exe -command ' + psscript).read()

		base.debugmsg(6, "response:", response)

	def create_macos_app_bundle(self, name, version, exesrc, icosrc):

		appspath = "~/Applications"
		if os.access("/Applications", os.W_OK):
			appspath = "/Applications"

		appspath = os.path.expanduser(appspath)

		# https://stackoverflow.com/questions/7404792/how-to-create-mac-application-bundle-for-python-script-via-python

		apppath = os.path.join(appspath, name + ".app")
		MacOSFolder = os.path.join(apppath, "Contents", "MacOS")
		base.ensuredir(MacOSFolder)

		# need to create the icon file:
		# https://stackoverflow.com/questions/646671/how-do-i-set-the-icon-for-my-applications-mac-os-x-app-bundle
		namelst = name.split()
		base.debugmsg(6, "namelst:", namelst)
		projname = "-".join(namelst).lower()
		base.debugmsg(6, "projname:", projname)
		signature = "RFS{0}".format(namelst[1].upper())
		base.debugmsg(6, "signature:", signature)

		ResourcesFolder = os.path.join(apppath, "Contents", "Resources")
		iconset = os.path.join(ResourcesFolder, projname + ".iconset")
		icnsfile = os.path.join(ResourcesFolder, projname + ".icns")
		base.ensuredir(iconset)

		# Normal screen icons
		base.debugmsg(6, "Normal screen icons")
		for size in [16, 32, 64, 128, 256, 512]:
			cmd = "sips -z {0} {0} {1} --out '{2}/icon_{0}x{0}.png'".format(size, icosrc, iconset)
			base.debugmsg(6, "cmd:", cmd)
			response = os.popen(cmd).read()
			base.debugmsg(6, "response:", response)

		# Retina display icons
		base.debugmsg(6, "Retina display icons")
		for size in [32, 64, 128, 256, 512, 1024]:
			cmd = "sips -z {0} {0} {1} --out '{2}/icon_{3}x{3}x2.png'".format(size, icosrc, iconset, int(size / 2))
			base.debugmsg(6, "cmd:", cmd)
			response = os.popen(cmd).read()
			base.debugmsg(6, "response:", response)

		# Make a multi-resolution Icon
		base.debugmsg(6, "Make a multi-resolution Icon")
		cmd = "iconutil -c icns -o '{0}' '{1}'".format(icnsfile, iconset)
		base.debugmsg(6, "cmd:", cmd)
		response = os.popen(cmd).read()
		base.debugmsg(6, "response:", response)

		#  create apppath + "/Contents/Info.plist"
		bundleName = name
		bundleIdentifier = "org.rfswarm." + projname

		# https://stackoverflow.com/questions/1596945/building-osx-app-bundle
		# Found 2 issues:
		# 	- <xml and <plist wasn't closed with > and xml was missing encoding
		# 	- APPL???? --> RFS<SIGNATURE_NAME>

		Infoplist = os.path.join(apppath, "Contents", "Info.plist")
		with open(Infoplist, "w") as f:
			f.write("""<?xml version="1.0" encoding="UTF-8"?>
			<!DOCTYPE plist PUBLIC "-//Apple//DTD PLIST 1.0//EN" "http://www.apple.com/DTDs/PropertyList-1.0.dtd">
			<plist version="1.0">
			<dict>
				<key>CFBundleDevelopmentRegion</key>
				<string>English</string>
				<key>CFBundleExecutable</key>
				<string>%s</string>
				<key>CFBundleGetInfoString</key>
				<string>%s</string>
				<key>CFBundleIconFile</key>
				<string>%s.icns</string>
				<key>CFBundleIdentifier</key>
				<string>%s</string>
				<key>CFBundleInfoDictionaryVersion</key>
				<string>6.0</string>
				<key>CFBundleName</key>
				<string>%s</string>
				<key>CFBundlePackageType</key>
				<string>APPL</string>
				<key>CFBundleShortVersionString</key>
				<string>%s</string>
				<key>CFBundleSignature</key>
				<string>%s</string>
				<key>CFBundleVersion</key>
				<string>%s</string>
				<key>NSAppleScriptEnabled</key>
				<string>YES</string>
				<key>NSMainNibFile</key>
				<string>MainMenu</string>
				<key>NSPrincipalClass</key>
				<string>NSApplication</string>
			</dict>
			</plist>
			""" % (projname, bundleName + " " + version, projname, bundleIdentifier, bundleName, version, signature, version))
			f.close()

		# create apppath + "/Contents/PkgInfo"
		PkgInfo = os.path.join(apppath, "Contents", "PkgInfo")
		with open(PkgInfo, "w") as f:
			f.write("APPL%s" % signature)
			f.close()

		# apppath + "/Contents/MacOS/main.py"
		execbundle = os.path.join(apppath, "Contents", "MacOS", projname)
		if os.path.exists(execbundle):
			os.remove(execbundle)
		os.symlink(exesrc, execbundle)

		# touch '/Applications/RFSwarm Reporter.app' to update .app icon
		cmd = "touch '{0}'".format(apppath)
		base.debugmsg(6, "cmd:", cmd)
		response = os.popen(cmd).read()
		base.debugmsg(6, "response:", response)

		# # Try re-registering your application with Launch Services:
		# # /System/Library/Frameworks/CoreServices.framework/Frameworks/LaunchServices.framework/Support/lsregister -f /Applications/MyTool.app
		# lsregister = "/System/Library/Frameworks/CoreServices.framework/Frameworks/LaunchServices.framework/Support/lsregister"
		# cmd = "{0} -f '{1}'".format(lsregister, apppath)
		# base.debugmsg(6, "cmd:", cmd)
		# response = os.popen(cmd).read()
		# base.debugmsg(6, "response:", response)

	def selectResults(self, resultsfile):
		base.debugmsg(5, "resultsfile:", resultsfile)

		if len(resultsfile) > 0:

			base.stop_db()

			base.config['Reporter']['Results'] = resultsfile

			tplfres = os.path.splitext(resultsfile)
			freport = "{}.report".format(tplfres[0])
			base.debugmsg(9, "freport:", freport)
			base.config['Reporter']['Report'] = freport

			filedir = os.path.dirname(resultsfile)
			base.debugmsg(9, "filedir:", filedir)
			parent = os.path.dirname(filedir)
			base.debugmsg(9, "parent:", parent)
			base.config['Reporter']['ResultDir'] = parent
			base.saveini()
			# base.open_results_db(base.config['Reporter']['Results'])
			base.start_db()

	def display_message(self, *mesage):
		if base.displaygui:
			msglst = []
			for msg in mesage:
				msglst.append(msg)
			msgout = " ".join(msglst)
			while base.gui is None and base.running:
				time.sleep(0.5)
			base.gui.updateStatus(msgout)
			base.debugmsg(1, msgout)
		else:
			msglst = []
			for msg in mesage:
				msglst.append(msg)
			msgout = " ".join(msglst)
			base.debugmsg(1, msgout)

	def export_xhtml(self):
		self.display_message("Generating XHTML Report")

		sections = base.report_get_order("TOP")
		base.debugmsg(5, "sections:", sections)
		sectionpct = 1 / (len(sections) + 1)
		base.debugmsg(5, "sectionpct:", sectionpct)

		if "XHTML" in self.cg_data:
			if "progress" in self.cg_data["XHTML"] and self.cg_data["XHTML"]["progress"] < 1:
				self.display_message("Waiting for previous XHTML report to finish")
				while self.cg_data["XHTML"]["progress"] < 1:
					time.sleep(0.5)
			del self.cg_data["XHTML"]
		self.cg_data["XHTML"] = {}
		self.cg_data["XHTML"]["progress"] = 0.0

		self.cg_data["XHTML"]["html"] = self.xhtml_base_doc()

		# set HTML Title
		title = self.cg_data["XHTML"]["html"].xpath("//head/title")[0]
		title.text = base.rs_setting_get_title()

		head = self.cg_data["XHTML"]["html"].xpath("//head")[0]

		style = etree.SubElement(head, 'style')
		highlightcolour = base.rs_setting_get_hcolour()
		fontname = base.rs_setting_get_font()
		fontsize = base.rs_setting_get_fontsize()

		styledata = ""
		styledata += "div { font-size: " + str(fontsize) + "px; font-family: \"" + fontname + "\"; }"
		styledata += ".center { text-align: center; }"
		styledata += ".title { font-size: 200%;}"
		styledata += ".subtitle { font-size: 150%;}"

		styledata += "table, th, td { border: 1px solid #ccc; border-collapse: collapse; }"
		styledata += "th { color: " + highlightcolour + "; }"

		# pre	{white-space: pre-wrap;}
		styledata += "pre	{white-space: pre-wrap;}"

		for i in range(6):
			styledata += "h" + str(i + 1) + "	{ color: " + highlightcolour + "; margin-left: " + str(i * 5) + "px; }"

		bodyindent = 30
		styledata += ".body { margin-left: " + str(bodyindent) + "px; }"
		#   margin-left: 20px;
		for i in range(6):
			styledata += ".TOC" + str(i + 1) + "	{margin-left: " + str(i * 10) + "px;}"

		style.text = styledata

		body = self.cg_data["XHTML"]["html"].xpath("//body")[0]

		# self.cg_data["XHTML"]["html"] =
		self.xhtml_add_sections(body, "TOP", sectionpct)

		base.debugmsg(5, "Report:", base.config['Reporter']['Report'])
		reportbase, reportext = os.path.splitext(base.config['Reporter']['Report'])

		outfile = "{}.html".format(reportbase)
		self.xhtml_save(outfile, self.cg_data["XHTML"]["html"])

		base.debugmsg(5, "outfile:", outfile)
		self.display_message("Saved XHTML Report:", outfile)

		self.cg_data["XHTML"]["progress"] = 1

	def export_pdf(self):
		base.debugmsg(5, "Not implimented yet.....")

	def export_word(self):
		self.display_message("Generating Word Report")

		sections = base.report_get_order("TOP")
		base.debugmsg(5, "sections:", sections)
		sectionpct = 1 / (len(sections) + 1)
		base.debugmsg(5, "sectionpct:", sectionpct)

		if "docx" in self.cg_data:
			if "progress" in self.cg_data["docx"] and self.cg_data["docx"]["progress"] < 1:
				self.display_message("Waiting for previous Word report to finish")
				while self.cg_data["docx"]["progress"] < 1:
					time.sleep(0.5)
			del self.cg_data["docx"]
		self.cg_data["docx"] = {}
		self.cg_data["docx"]["progress"] = 0.0

		self.cg_data["docx"]["document"] = Document()

		self.docx_configure_style()

		self.docx_add_sections("TOP", sectionpct)

		base.debugmsg(5, "Report:", base.config['Reporter']['Report'])
		reportbase, reportext = os.path.splitext(base.config['Reporter']['Report'])
		outfile = "{}.docx".format(reportbase)
		self.cg_data["docx"]["document"].save(outfile)

		base.debugmsg(5, "outfile:", outfile)
		self.display_message("Saved Word Report:", outfile)

		self.cg_data["docx"]["progress"] = 1

	def export_writer(self):
		base.debugmsg(5, "Not implimented yet.....")

	def export_excel(self):
		self.display_message("Generating Excel Report")

		sections = base.report_get_order("TOP")
		base.debugmsg(5, "sections:", sections)
		sectionpct = 1 / (len(sections) + 1)
		base.debugmsg(5, "sectionpct:", sectionpct)

		if "xlsx" in self.cg_data:
			if "progress" in self.cg_data["xlsx"] and self.cg_data["xlsx"]["progress"] < 1:
				self.display_message("Waiting for previous Excel report to finish")
				while self.cg_data["xlsx"]["progress"] < 1:
					time.sleep(0.5)
			del self.cg_data["xlsx"]
		self.cg_data["xlsx"] = {}
		self.cg_data["xlsx"]["progress"] = 0.0

		self.cg_data["xlsx"]["Workbook"] = openpyxl.Workbook()

		self.xlsx_configure_style()

		self.xlsx_add_sections("TOP", sectionpct)

		self.cg_data["xlsx"]["Workbook"].active = self.cg_data["xlsx"]["Workbook"].worksheets[0]

		base.debugmsg(5, "Report:", base.config['Reporter']['Report'])
		reportbase, reportext = os.path.splitext(base.config['Reporter']['Report'])
		outfile = "{}.xlsx".format(reportbase)
		self.cg_data["xlsx"]["Workbook"].save(outfile)

		base.debugmsg(5, "outfile:", outfile)
		self.display_message("Saved Excel Report:", outfile)

		self.cg_data["xlsx"]["progress"] = 1

	def export_calc(self):
		base.debugmsg(5, "Not implimented yet.....")

	#
	# 	XHTML
	#

	def xhtml_save(self, filename, html):
		result = etree.tostring(
			html,
			xml_declaration=True,
			doctype='<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">',
			encoding='utf-8',
			standalone=False,
			with_tail=False,
			method='html',
			pretty_print=True)

		with open(filename, "wb") as f:
			f.write(result)

	def xhtml_base_doc(self):
		# https://www.w3schools.com/html/html_xhtml.asp

		# <!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN"
		# "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
		# <html xmlns="http://www.w3.org/1999/xhtml">
		# <head>
		#   <title>Title of document</title>
		# </head>
		# <body>
		#
		#   some content here...
		#
		# </body>
		# </html>

		M = ElementMaker(namespace=None, nsmap={None: "http://www.w3.org/1999/xhtml"})
		# html = M.html(E.head(E.title("Test page")), E.body(E.p("Hello world")))
		html = M.html(E.head(E.title("Test page"), E.meta(charset="utf-8")), E.body())

		return html

	def xhtml_add_sections(self, parent, id, sectionpct):
		base.debugmsg(8, "id:", id, "	sectionpct:", sectionpct)
		# SubElement(_parent, _tag, attrib=None, nsmap=None, **_extra)
		# thiselmt = E.SubElement(parent, 'div', attrib={'id':id})

		base.debugmsg(9, "parent:", parent)

		sections = base.report_get_order(id)
		base.debugmsg(9, "sections:", sections)

		thiselmt = etree.SubElement(parent, 'div')
		nextparent = thiselmt
		if id == "TOP":
			thiselmt.set("id", "TitlePage")
			nextparent = parent

			#
			# Title
			#
			maintitle = etree.SubElement(thiselmt, 'div')
			maintitle.set("class", "title center")
			maintitle.text = base.rs_setting_get_title()

			#
			# Logo
			#
			base.debugmsg(5, "showtlogo:", base.rs_setting_get_int("showtlogo"))
			if base.rs_setting_get_int("showtlogo"):
				imgtitle = etree.SubElement(thiselmt, 'div')
				imgtitle.set("class", "center")

				tlogo = base.rs_setting_get_file("tlogo")
				base.debugmsg(5, "tlogo:", tlogo)
				if len(tlogo) > 0:
					self.xhtml_sections_fileimg(imgtitle, id, tlogo)

			#
			# Execution Date range
			#
			subtitle = etree.SubElement(thiselmt, 'div')
			subtitle.set("class", "subtitle center")
			execdr = ""
			if base.rs_setting_get_int("showstarttime"):
				iST = base.rs_setting_get_starttime()
				fSD = "{}".format(base.report_formatdate(iST))
				fST = "{}".format(base.report_formattime(iST))

				execdr = "{} {}".format(fSD, fST)

			if base.rs_setting_get_int("showendtime"):
				iET = base.rs_setting_get_endtime()
				fED = "{}".format(base.report_formatdate(iET))
				fET = "{}".format(base.report_formattime(iET))

				if not base.rs_setting_get_int("showstarttime"):
					execdr = "{} {}".format(fED, fET)
				else:
					if fSD == fED:
						execdr = "{} - {}".format(execdr, fET)
					else:
						execdr = "{} - {} {}".format(execdr, fED, fET)

			subtitle.text = execdr

		else:
			thiselmt.set("id", id)
			newsectionpct = 1 / (len(sections) + 1)
			sectionpct = newsectionpct * sectionpct
			base.debugmsg(9, "sectionpct:", sectionpct)
			self.xhtml_sections_addheading(thiselmt, id)

			stype = base.report_item_get_type(id)
			base.debugmsg(9, "stype:", stype)
			if stype == "contents":
				self.xhtml_sections_contents(thiselmt, id)
			if stype == "note":
				self.xhtml_sections_note(thiselmt, id)
			if stype == "graph":
				self.xhtml_sections_graph(thiselmt, id)
			if stype == "table":
				self.xhtml_sections_table(thiselmt, id)
			if stype == "errors":
				self.xhtml_sections_errors(thiselmt, id)

		self.cg_data["XHTML"]["progress"] += sectionpct
		self.display_message("Generating XHTML Report {}%".format(int(round(self.cg_data["XHTML"]["progress"] * 100, 0))))

		if len(sections) > 0:
			for sect in sections:
				self.xhtml_add_sections(nextparent, sect, sectionpct)

	def xhtml_sections_addheading(self, elmt, id):
		base.debugmsg(8, "id:", id)
		level = base.report_sect_level(id)
		base.debugmsg(9, "level:", level)
		number = base.report_sect_number(id)
		base.debugmsg(9, "number:", number)
		name = base.report_item_get_name(id)
		base.debugmsg(9, "name:", name)

		tag = "h{}".format(level)

		h = etree.SubElement(elmt, tag)
		h.text = "{} {}".format(number, name)
		# a = etree.SubElement(h, 'a')
		# a.text = "{} {}".format(number, name)
		# a.set('href', "#{}".format(id))

	def xhtml_sections_fileimg(self, elmt, id, imgfile):
		base.debugmsg(5, "id:", id, "	imgfile:", imgfile)
		# img.set("src", imgfile)

		oimg = Image.open(imgfile)
		base.debugmsg(9, "oimg:", oimg)
		self.xhtml_sections_embedimg(elmt, id, oimg)

	def xhtml_sections_embedimg(self, elmt, id, oimg):
		base.debugmsg(8, "id:", id, "	oimg:", oimg)

		img = etree.SubElement(elmt, 'img')

		img.set("id", id + "_img")
		img.set("imgid", id)
		img.text = ""

		# <img src="data:image/png;base64,
		base.debugmsg(9, "oimg format:", oimg.format)
		# base.debugmsg(5, "oimg info:", oimg.info)

		# <img src="data:image/png;base64,
		srcdata = "data:image/" + oimg.format.lower() + ";base64,"

		# https://stackoverflow.com/questions/48229318/how-to-convert-image-pil-into-base64-without-saving
		# srcdata += str(base64.b64encode(oimg.tobytes()))
		# srcdata += base64.b64encode(oimg.tobytes()) # bytes
		# srcdata += base64.b64encode(oimg.tobytes()).decode()

		buffered = BytesIO()
		oimg.save(buffered, format=oimg.format)
		buffered.seek(0)
		img_byte = buffered.getvalue()

		srcdata += base64.b64encode(img_byte).decode()

		base.debugmsg(9, "srcdata:", srcdata)
		img.set("src", srcdata)

	def xhtml_sections_contents(self, elmt, id):
		base.debugmsg(8, "id:", id)
		mode = base.rt_contents_get_mode(id)
		level = base.rt_contents_get_level(id)

		base.debugmsg(9, "mode:", mode, "	level:", level)
		fmode = None
		if mode == "Table Of Contents":
			fmode = None
		if mode == "Table of Graphs":
			fmode = "graph"
		if mode == "Table Of Tables":
			fmode = "table"

		# tbl = etree.SubElement(elmt, 'table')
		body = etree.SubElement(elmt, 'div')
		body.set("class", "body")

		self.xhtml_sections_contents_row(body, "TOP", 1, fmode, level)
		# self.xhtml_sections_contents_row(tbl, "TOP", 1, fmode, level)

	def xhtml_sections_contents_row(self, elmt, id, rownum, fmode, flevel):
		base.debugmsg(8, "id:", id, "	rownum:", rownum, "	fmode:", fmode, "	flevel:", flevel)
		display = True

		level = base.report_sect_level(id)
		if id == "TOP":
			display = False
			level = 0
		base.debugmsg(9, "level:", level)

		if display and fmode is not None:
			display = False
			type = base.report_item_get_type(id)
			if fmode == type:
				display = True

		if display and level > flevel:
			display = False

		nextrow = rownum
		if display:
			type = base.report_item_get_type(id)
			titlenum = base.report_sect_number(id)
			titlename = base.report_item_get_name(id)
			# titlelevel = base.report_sect_level(id)

			p = etree.SubElement(elmt, 'p')
			a = etree.SubElement(p, 'a')
			a.set("class", "TOC{}".format(level))
			a.text = "{} {}".format(titlenum, titlename)
			if fmode is None:
				idpre = "toc"
			else:
				idpre = fmode
			a.set('id', "{}_{}".format(idpre, id))
			a.set('href', "#{}".format(id))

			nextrow = rownum + 1
		base.debugmsg(9, "nextrow:", nextrow)
		if level < flevel:
			children = base.report_get_order(id)
			for child in children:
				nextrow = self.xhtml_sections_contents_row(elmt, child, nextrow, fmode, flevel)
		return nextrow

	def xhtml_sections_note(self, elmt, id):
		base.debugmsg(8, "id:", id)
		notebody = base.rt_note_get(id)
		notebody = notebody.replace("\r\n", "\n")
		notebody = notebody.replace("\r", "\n")
		notelist = notebody.split("\n")
		base.debugmsg(9, "notebody:", notebody)
		body = etree.SubElement(elmt, 'div')
		body.set("class", "body")
		for line in notelist:
			p = etree.SubElement(body, "p")
			p.text = line

	def xhtml_sections_graph(self, elmt, id):
		base.debugmsg(8, "id:", id)
		pid, idl, idr = base.rt_graph_LR_Ids(id)
		base.debugmsg(5, "pid:", pid, "	idl:", idl, "	idr:", idr)

		body = etree.SubElement(elmt, 'div')
		body.set("class", "body")

		axisenl = base.rt_graph_get_axisen(idl)
		axisenr = base.rt_graph_get_axisen(idr)

		datatypel = base.rt_graph_get_dt(idl)
		datatyper = base.rt_graph_get_dt(idr)

		tz = zoneinfo.ZoneInfo(base.rs_setting_get_timezone())

		if datatypel == "SQL":
			sqll = base.rt_graph_get_sql(idl)
		else:
			sqll = base.rt_graph_generate_sql(idl)

		if datatyper == "SQL":
			sqlr = base.rt_graph_get_sql(idr)
		else:
			sqlr = base.rt_graph_generate_sql(idr)

		gphdpi = 72
		# gphdpi = 100
		fig = Figure(dpi=gphdpi)
		axisl = fig.add_subplot(1, 1, 1)
		axisl.grid(True, 'major', 'x')
		axisr = axisl.twinx()

		fig.autofmt_xdate(bottom=0.2, rotation=30, ha='right')

		canvas = FigureCanvas(fig)

		# https://stackoverflow.com/questions/57316491/how-to-convert-matplotlib-figure-to-pil-image-object-without-saving-image
		axisl.tick_params(labelleft=False, length=0)
		axisr.tick_params(labelright=False, length=0)

		try:
			canvas.draw()
		except Exception as e:
			base.debugmsg(3, "canvas.draw() Exception:", e)
		fig.set_tight_layout(True)

		dodraw = False
		graphdata = {}

		if (sqll is not None and len(sqll.strip()) > 0) or (sqlr is not None and len(sqlr.strip()) > 0):

			if sqll is not None and len(sqll.strip()) > 0 and axisenl > 0:
				base.debugmsg(7, "sqll:", sqll)
				key = "{}_{}".format(idl, base.report_item_get_changed(idl))
				base.dbqueue["Read"].append({"SQL": sqll, "KEY": key})
				while key not in base.dbqueue["ReadResult"]:
					time.sleep(0.1)

				gdata = base.dbqueue["ReadResult"][key]

				if datatypel == "Plan":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_plan(idl, gdata)
				if datatypel == "Monitoring":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_monitoring(idl, gdata)

				base.debugmsg(9, "gdata:", gdata)

				for row in gdata:
					base.debugmsg(9, "row:", row)
					if 'Name' in row:
						name = row['Name']
						base.debugmsg(9, "name:", name)
						if name not in graphdata:
							graphdata[name] = {}

							colour = base.named_colour(name)
							base.debugmsg(8, "name:", name, "	colour:", colour)
							graphdata[name]["Colour"] = colour
							graphdata[name]["Axis"] = "axisL"
							# self.contentdata[id]["graphdata"][name]["Time"] = []
							graphdata[name]["objTime"] = []
							graphdata[name]["Values"] = []

						graphdata[name]["objTime"].append(datetime.fromtimestamp(row["Time"], tz))
						graphdata[name]["Values"].append(base.rt_graph_floatval(row["Value"]))
					else:
						break

			if sqlr is not None and len(sqlr.strip()) > 0 and axisenr > 0:
				base.debugmsg(7, "sqlr:", sqlr)
				key = "{}_{}".format(idr, base.report_item_get_changed(idr))
				base.dbqueue["Read"].append({"SQL": sqlr, "KEY": key})
				while key not in base.dbqueue["ReadResult"]:
					time.sleep(0.1)

				gdata = base.dbqueue["ReadResult"][key]

				if datatyper == "Plan":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_plan(idr, gdata)
				if datatyper == "Monitoring":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_monitoring(idr, gdata)

				base.debugmsg(9, "gdata:", gdata)

				for row in gdata:
					base.debugmsg(9, "row:", row)
					if 'Name' in row:
						name = row['Name']
						base.debugmsg(9, "name:", name)
						if name not in graphdata:
							graphdata[name] = {}

							colour = base.named_colour(name)
							base.debugmsg(8, "name:", name, "	colour:", colour)
							graphdata[name]["Colour"] = colour
							graphdata[name]["Axis"] = "axisR"
							# self.contentdata[id]["graphdata"][name]["Time"] = []
							graphdata[name]["objTime"] = []
							graphdata[name]["Values"] = []

						graphdata[name]["objTime"].append(datetime.fromtimestamp(row["Time"], tz))
						graphdata[name]["Values"].append(base.rt_graph_floatval(row["Value"]))
					else:
						break

			base.debugmsg(9, "graphdata:", graphdata)

			for name in graphdata:
				base.debugmsg(7, "name:", name)

				axis = "axisL"
				if "Axis" in graphdata[name]:
					axis = graphdata[name]["Axis"]

				if len(graphdata[name]["Values"]) > 1 and len(graphdata[name]["Values"]) == len(graphdata[name]["objTime"]):
					try:
						if axis == "axisL":
							axisl.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name)
						elif axis == "axisR":
							axisr.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name)
						dodraw = True
					except Exception as e:
						base.debugmsg(7, "axis.plot() Exception:", e)

				if len(graphdata[name]["Values"]) == 1 and len(graphdata[name]["Values"]) == len(graphdata[name]["objTime"]):
					try:
						if axis == "axisL":
							axisl.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name, marker='o')
						elif axis == "axisR":
							axisr.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name, marker='o')
						dodraw = True
					except Exception as e:
						base.debugmsg(7, "axis.plot() Exception:", e)

			if dodraw:

				# Left axis Limits
				if axisenl > 0:
					axisl.grid(True, 'major', 'y')
					axisl.tick_params(labelleft=True, length=5)

					SMetric = "Other"
					if datatypel == "Metric":
						SMetric = base.rt_table_get_sm(idl)
					base.debugmsg(8, "SMetric:", SMetric)
					if SMetric in ["Load", "CPU", "MEM", "NET"]:
						axisl.set_ylim(0, 100)
					else:
						axisl.set_ylim(0)

				# Right axis Limits
				if axisenr > 0:
					axisr.grid(True, 'major', 'y')
					axisr.tick_params(labelright=True, length=5)

					SMetric = "Other"
					if datatyper == "Metric":
						SMetric = base.rt_table_get_sm(idr)
					base.debugmsg(8, "SMetric:", SMetric)
					if SMetric in ["Load", "CPU", "MEM", "NET"]:
						axisr.set_ylim(0, 100)
					else:
						axisr.set_ylim(0)

				fig.set_tight_layout(True)
				fig.autofmt_xdate(bottom=0.2, rotation=30, ha='right')
				try:
					canvas.draw()
				except Exception as e:
					base.debugmsg(3, "canvas.draw() Exception:", e)

				buf = BytesIO()
				fig.savefig(buf)
				buf.seek(0)
				oimg = Image.open(buf)
				self.xhtml_sections_embedimg(body, id, oimg)

	def xhtml_sections_table(self, elmt, id):
		base.debugmsg(8, "id:", id)
		body = etree.SubElement(elmt, 'div')
		body.set("class", "body")
		tbl = etree.SubElement(body, 'table')

		datatype = base.rt_table_get_dt(id)
		if datatype == "SQL":
			sql = base.rt_table_get_sql(id)
		else:
			sql = base.rt_table_generate_sql(id)
		colours = base.rt_table_get_colours(id)

		if sql is not None and len(sql.strip()) > 0:
			base.debugmsg(9, "sql:", sql)
			key = "{}_{}".format(id, base.report_item_get_changed(id))
			base.dbqueue["Read"].append({"SQL": sql, "KEY": key})
			while key not in base.dbqueue["ReadResult"]:
				time.sleep(0.1)

			tdata = base.dbqueue["ReadResult"][key]
			if datatype == "Plan":
				base.debugmsg(9, "tdata before:", tdata)
				tdata = base.table_postprocess_data_plan(id, tdata)
			if datatype == "Monitoring":
				base.debugmsg(9, "tdata before:", tdata)
				tdata = base.table_postprocess_data_monitoring(id, tdata)
			base.debugmsg(9, "tdata:", tdata)

			if len(tdata) > 0:
				# table headers
				cols = list(tdata[0].keys())
				base.debugmsg(8, "cols:", cols)
				tr = etree.SubElement(tbl, 'tr')
				if colours:
					th = etree.SubElement(tr, 'th')
				for col in cols:
					if col not in ["Colour"]:
						show = base.report_item_get_bool_def1(id, base.rt_table_ini_colname(f"{col} Show"))
						if show:
							dispname = base.rt_table_get_colname(id, col)
							th = etree.SubElement(tr, 'th')
							th.text = dispname.strip()

				# table rows
				for row in tdata:
					vals = list(row.values())
					base.debugmsg(8, "vals:", vals)
					tr = etree.SubElement(tbl, 'tr')
					if colours:

						base.debugmsg(9, "row:", row)
						if "Colour" in row:
							label = row["Colour"]
						else:
							label = row[cols[0]]
						base.debugmsg(9, "label:", label)
						colour = base.named_colour(label)
						base.debugmsg(9, "colour:", colour)

						# <td style="background-color:#8888ff; color:#8888ff;">_</td>
						td = etree.SubElement(tr, 'td')
						td.text = "_"
						rstyle = "background-color:{}; color:{};".format(colour, colour)
						td.set("style", rstyle)

					# for val in vals:
					for col in cols:
						if col not in ["Colour"]:
							show = base.report_item_get_bool_def1(id, base.rt_table_ini_colname(f"{col} Show"))
							if show:
								td = etree.SubElement(tr, 'td')
								val = str(row[col]).strip()
								val = base.illegal_xml_chars_re.sub('', val)
								base.debugmsg(8, "val:", val)
								td.text = val

	def xhtml_sections_errors(self, elmt, id):
		base.debugmsg(8, "id:", id)
		body = etree.SubElement(elmt, 'div')
		body.set("class", "body")
		tbl = etree.SubElement(body, 'table')

		showimages = base.rt_errors_get_images(id)
		base.debugmsg(5, "showimages:", showimages)
		grouprn = base.rt_errors_get_group_rn(id)
		base.debugmsg(5, "grouprn:", grouprn)
		groupet = base.rt_errors_get_group_et(id)
		base.debugmsg(5, "groupet:", groupet)

		lbl_Result = base.rt_errors_get_label(id, "lbl_Result")
		lbl_Test = base.rt_errors_get_label(id, "lbl_Test")
		lbl_Script = base.rt_errors_get_label(id, "lbl_Script")
		lbl_Error = base.rt_errors_get_label(id, "lbl_Error")
		lbl_Count = base.rt_errors_get_label(id, "lbl_Count")
		lbl_Screenshot = base.rt_errors_get_label(id, "lbl_Screenshot")
		lbl_NoScreenshot = base.rt_errors_get_label(id, "lbl_NoScreenshot")

		pctalike = 0.80
		base.debugmsg(5, "pctalike:", pctalike)

		base.rt_errors_get_data(id)

		grpdata = {}
		if grouprn or groupet:
			grpdata = {}
			grpdata["resultnames"] = {}
			grpdata["errortexts"] = {}

			keys = list(base.reportdata[id].keys())
			for key in keys:
				base.debugmsg(5, "key:", key)
				if key != "key":
					rdata = base.reportdata[id][key]

					if grouprn:
						result_name = rdata['result_name']
						matches = difflib.get_close_matches(result_name, list(grpdata["resultnames"].keys()), cutoff=pctalike)
						base.debugmsg(5, "matches:", matches)
						if len(matches) > 0:
							result_name = matches[0]
							basekey = grpdata["resultnames"][result_name]["keys"][0]
							base.debugmsg(5, "basekey:", basekey)
							grpdata["resultnames"][result_name]["keys"].append(key)
						else:
							grpdata["resultnames"][result_name] = {}
							grpdata["resultnames"][result_name]["keys"] = []
							grpdata["resultnames"][result_name]["keys"].append(key)
							grpdata["resultnames"][result_name]["errortexts"] = {}

						if groupet and 'error' in rdata:
							errortext = rdata['error']
							# errortext_sub = errortext.split(r'\n')[0]
							errortext_sub = errortext.splitlines()[0]
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							matcheset = difflib.get_close_matches(errortext_sub, list(grpdata["resultnames"][result_name]["errortexts"].keys()), cutoff=pctalike)
							base.debugmsg(5, "matcheset:", matcheset)
							if len(matcheset) > 0:
								errortext = matcheset[0]
								errortext_sub = errortext.splitlines()[0]
								baseid = grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"][0]
								base.debugmsg(5, "baseid:", baseid)
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"].append(key)
							else:
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub] = {}
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"] = []
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"].append(key)

					if groupet and 'error' in rdata:
						errortext = rdata['error']
						errortext_sub = errortext.splitlines()[0]
						base.debugmsg(5, "errortext_sub:", errortext_sub)
						matches = difflib.get_close_matches(errortext_sub, list(grpdata["errortexts"].keys()), cutoff=pctalike)
						base.debugmsg(5, "matches:", matches)
						if len(matches) > 0:
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							errortext = matches[0]
							base.debugmsg(5, "errortext:", errortext)
							baseid = grpdata["errortexts"][errortext]["keys"][0]
							base.debugmsg(5, "baseid:", baseid)
							grpdata["errortexts"][errortext]["keys"].append(key)
						else:
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							grpdata["errortexts"][errortext_sub] = {}
							grpdata["errortexts"][errortext_sub]["keys"] = []
							grpdata["errortexts"][errortext_sub]["keys"].append(key)

			resultnames = grpdata["resultnames"]
			base.debugmsg(5, "resultnames:", resultnames)
			errortexts = grpdata["errortexts"]
			base.debugmsg(5, "errortexts:", errortexts)

		if grouprn:
			for result_name in list(grpdata["resultnames"].keys()):
				basekey = grpdata["resultnames"][result_name]["keys"][0]
				base.debugmsg(5, "basekey:", basekey)
				rdata = base.reportdata[id][basekey]
				tr = etree.SubElement(tbl, 'tr')

				th = etree.SubElement(tr, 'th')
				th.text = "{}:".format(lbl_Result)
				td = etree.SubElement(tr, 'td')
				td.text = result_name

				th = etree.SubElement(tr, 'th')
				th.text = "{}:".format(lbl_Test)
				td = etree.SubElement(tr, 'td')
				td.text = rdata['test_name']

				th = etree.SubElement(tr, 'th')
				th.text = "{}:".format(lbl_Script)
				td = etree.SubElement(tr, 'td')
				td.text = rdata['script']

				th = etree.SubElement(tr, 'th')
				th.text = "{}:".format(lbl_Count)
				count = len(grpdata["resultnames"][result_name]["keys"])
				base.debugmsg(5, "count:", count)
				td = etree.SubElement(tr, 'td')
				td.text = str(count)

				if groupet:
					for errortext in list(grpdata["resultnames"][result_name]["errortexts"].keys()):
						basekey = grpdata["resultnames"][result_name]["errortexts"][errortext]["keys"][0]
						base.debugmsg(5, "basekey:", basekey)
						rdata = base.reportdata[id][basekey]

						tr = etree.SubElement(tbl, 'tr')

						th = etree.SubElement(tr, 'th')
						th.text = "{}:".format(lbl_Error)
						td = etree.SubElement(tr, 'td')
						pre = etree.SubElement(td, 'pre')
						if 'error' in rdata:
							pre.text = rdata['error']
							td.set('colspan', '5')

							th = etree.SubElement(tr, 'th')
							th.text = "{}:".format(lbl_Count)
							td = etree.SubElement(tr, 'td')
							count = len(grpdata["resultnames"][result_name]["errortexts"][errortext]["keys"])
							base.debugmsg(5, "count:", count)
							td.text = str(count)

							if showimages:
								tr = etree.SubElement(tbl, 'tr')

								th = etree.SubElement(tr, 'th')
								th.text = "{}:".format(lbl_Screenshot)

								td = etree.SubElement(tr, 'td')
								td.set('colspan', '7')
								if 'image_file' in rdata:
									# td.text = rdata['image_file']
									oimg = Image.open(rdata['image_file'])
									self.xhtml_sections_embedimg(td, basekey, oimg)
								else:
									td.text = lbl_NoScreenshot

				else:
					for keyi in grpdata["resultnames"][result_name]["keys"]:
						rdata = base.reportdata[id][keyi]

						tr = etree.SubElement(tbl, 'tr')

						th = etree.SubElement(tr, 'th')
						th.text = "{}:".format(lbl_Error)

						td = etree.SubElement(tr, 'td')
						pre = etree.SubElement(td, 'pre')
						if 'error' in rdata:
							pre.text = rdata['error']
							td.set('colspan', '7')

							if showimages:
								tr = etree.SubElement(tbl, 'tr')

								th = etree.SubElement(tr, 'th')
								th.text = "{}:".format(lbl_Screenshot)

								td = etree.SubElement(tr, 'td')
								td.set('colspan', '7')
								if 'image_file' in rdata:
									# td.text = rdata['image_file']
									oimg = Image.open(rdata['image_file'])
									self.xhtml_sections_embedimg(td, keyi, oimg)
								else:
									td.text = lbl_NoScreenshot

		if groupet and not grouprn:
			for errortext in list(grpdata["errortexts"].keys()):
				basekey = grpdata["errortexts"][errortext]["keys"][0]
				base.debugmsg(5, "basekey:", basekey)
				rdata = base.reportdata[id][basekey]

				tr = etree.SubElement(tbl, 'tr')

				th = etree.SubElement(tr, 'th')
				th.text = "{}:".format(lbl_Error)

				td = etree.SubElement(tr, 'td')
				pre = etree.SubElement(td, 'pre')
				if 'error' in rdata:
					pre.text = rdata['error']
					td.set('colspan', '5')

					th = etree.SubElement(tr, 'th')
					th.text = "{}:".format(lbl_Count)
					count = len(grpdata["errortexts"][errortext]["keys"])
					base.debugmsg(5, "count:", count)
					td = etree.SubElement(tr, 'td')
					td.text = str(count)

					if showimages:
						tr = etree.SubElement(tbl, 'tr')

						th = etree.SubElement(tr, 'th')
						th.text = "{}:".format(lbl_Screenshot)

						td = etree.SubElement(tr, 'td')
						td.set('colspan', '7')
						if 'image_file' in rdata:
							# td.text = rdata['image_file']
							oimg = Image.open(rdata['image_file'])
							self.xhtml_sections_embedimg(td, basekey, oimg)
						else:
							td.text = lbl_NoScreenshot

		if not grouprn and not groupet:
			keys = list(base.reportdata[id].keys())
			for key in keys:
				base.debugmsg(5, "key:", key)
				if key != "key":
					rdata = base.reportdata[id][key]

					tr = etree.SubElement(tbl, 'tr')

					th = etree.SubElement(tr, 'th')
					th.text = "{}:".format(lbl_Result)
					td = etree.SubElement(tr, 'td')
					td.text = rdata['result_name']

					th = etree.SubElement(tr, 'th')
					th.text = "{}:".format(lbl_Test)
					td = etree.SubElement(tr, 'td')
					td.text = rdata['test_name']

					th = etree.SubElement(tr, 'th')
					th.text = "{}:".format(lbl_Script)
					td = etree.SubElement(tr, 'td')
					td.text = rdata['script']

					tr = etree.SubElement(tbl, 'tr')

					th = etree.SubElement(tr, 'th')
					th.text = "{}:".format(lbl_Error)
					td = etree.SubElement(tr, 'td')
					pre = etree.SubElement(td, 'pre')
					if 'error' in rdata:
						pre.text = rdata['error']
						td.set('colspan', '5')

						if showimages:
							tr = etree.SubElement(tbl, 'tr')

							th = etree.SubElement(tr, 'th')
							th.text = "{}:".format(lbl_Screenshot)

							td = etree.SubElement(tr, 'td')
							td.set('colspan', '5')
							if 'image_file' in rdata:
								# td.text = rdata['image_file']
								oimg = Image.open(rdata['image_file'])
								self.xhtml_sections_embedimg(td, key, oimg)
							else:
								td.text = lbl_NoScreenshot

	#
	# 	MS Word
	#
	# https://python-docx.readthedocs.io/en/latest/
	# https://github.com/python-openxml/python-docx

	def docx_configure_style(self):
		# set up document styles for this report
		highlightcolour = base.rs_setting_get_hcolour().replace("#", "")
		fontname = base.rs_setting_get_font()
		fontsize = base.rs_setting_get_fontsize()

		# rgb_basecolour = RGBColor.from_string('000000')
		rgb_highlightcolour = RGBColor.from_string(highlightcolour)
		base.debugmsg(5, "rgb_highlightcolour:", rgb_highlightcolour)

		base.debugmsg(5, "fontname:", fontname, "	fontsize:", fontsize, "	highlightcolour:", highlightcolour)

		# styles = self.cg_data["docx"]["document"].styles
		# for style in styles:
		# 	print(style.name)

		# Update Normal
		style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize)
		base.debugmsg(5, "style.paragraph_format.left_indent:", style.paragraph_format.left_indent)
		style.paragraph_format.left_indent = Cm(0.5)
		base.debugmsg(5, "style.paragraph_format.left_indent:", style.paragraph_format.left_indent)

		# Update Cover Title
		self.cg_data["docx"]["document"].styles.add_style('Cover Title', WD_STYLE_TYPE.PARAGRAPH)
		style = self.cg_data["docx"]["document"].styles['Cover Title']
		style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.size = Pt(int(fontsize * 2))
		style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		style.paragraph_format.left_indent = Cm(0)

		# Update Subtitle
		self.cg_data["docx"]["document"].styles.add_style('Cover Subtitle', WD_STYLE_TYPE.PARAGRAPH)
		style = self.cg_data["docx"]["document"].styles['Cover Subtitle']
		style.base_style = self.cg_data["docx"]["document"].styles['Cover Title']
		style.font.size = Pt(int(fontsize * 1.5))

		sizeup = int(fontsize * 0.1)
		if sizeup < 1:
			sizeup = int(1)
		base.debugmsg(8, "sizeup:", sizeup)

		# Update Heading 1
		style = self.cg_data["docx"]["document"].styles['Heading 1']
		# style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize + (6 * sizeup))
		style.font.italic = False
		style.font.color.rgb = rgb_highlightcolour
		# style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		style.paragraph_format.page_break_before = True
		style.paragraph_format.keep_with_next = True
		style.paragraph_format.left_indent = Cm(0)

		# Update Heading 2
		style = self.cg_data["docx"]["document"].styles['Heading 2']
		# style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize + (5 * sizeup))
		style.font.italic = False
		style.font.color.rgb = rgb_highlightcolour
		style.paragraph_format.page_break_before = False
		style.paragraph_format.keep_with_next = True
		style.paragraph_format.left_indent = Cm(0)

		# Update Heading 3
		style = self.cg_data["docx"]["document"].styles['Heading 3']
		# style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize + (4 * sizeup))
		style.font.italic = False
		style.font.color.rgb = rgb_highlightcolour
		style.paragraph_format.page_break_before = False
		style.paragraph_format.keep_with_next = True
		style.paragraph_format.left_indent = Cm(0)

		# Update Heading 4
		style = self.cg_data["docx"]["document"].styles['Heading 4']
		# style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize + (3 * sizeup))
		style.font.italic = False
		style.font.color.rgb = rgb_highlightcolour
		style.paragraph_format.page_break_before = False
		style.paragraph_format.keep_with_next = True
		style.paragraph_format.left_indent = Cm(0)

		# Update Heading 5
		style = self.cg_data["docx"]["document"].styles['Heading 5']
		# style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize + (2 * sizeup))
		style.font.italic = False
		style.font.color.rgb = rgb_highlightcolour
		style.paragraph_format.page_break_before = False
		style.paragraph_format.keep_with_next = True
		style.paragraph_format.left_indent = Cm(0)

		# Update Heading 6
		style = self.cg_data["docx"]["document"].styles['Heading 6']
		# style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.name = fontname
		style.font.size = Pt(fontsize + (1 * sizeup))
		style.font.italic = False
		style.font.color.rgb = rgb_highlightcolour
		style.paragraph_format.page_break_before = False
		style.paragraph_format.keep_with_next = True
		style.paragraph_format.left_indent = Cm(0)

		# Update Table Heading?

		# Table Cell
		self.cg_data["docx"]["document"].styles.add_style('Table Cell', WD_STYLE_TYPE.PARAGRAPH)
		style = self.cg_data["docx"]["document"].styles['Table Cell']
		style.base_style = self.cg_data["docx"]["document"].styles['Normal']
		style.font.size = Pt(int(fontsize * 0.8))
		style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		# style.paragraph_format.left_indent = Cm(0)
		style.paragraph_format.left_indent = Cm(-0.15)
		# style.paragraph_format.right_indent = Cm(0)
		style.paragraph_format.right_indent = Cm(-0.15)
		style.paragraph_format.space_after = Cm(0.1)

		# Table Header
		self.cg_data["docx"]["document"].styles.add_style('Table Header', WD_STYLE_TYPE.PARAGRAPH)
		style = self.cg_data["docx"]["document"].styles['Table Header']
		style.base_style = self.cg_data["docx"]["document"].styles['Table Cell']
		style.font.bold = True
		style.font.color.rgb = rgb_highlightcolour

	def docx_add_sections(self, id, sectionpct):
		base.debugmsg(5, "id:", id, "	sectionpct:", sectionpct)

		sections = base.report_get_order(id)
		base.debugmsg(5, "sections:", sections)

		document = self.cg_data["docx"]["document"]

		if id == "TOP":

			#
			# Title
			#
			titletxt = base.rs_setting_get_title()
			# document.add_heading(titletxt, 0)
			document.add_paragraph("", style='Cover Title')
			document.add_paragraph(titletxt, style='Cover Title')
			document.add_paragraph("", style='Cover Title')

			#
			# Logo
			#
			base.debugmsg(5, "showtlogo:", base.rs_setting_get_int("showtlogo"))
			if base.rs_setting_get_int("showtlogo"):

				tlogo = base.rs_setting_get_file("tlogo")
				base.debugmsg(5, "tlogo:", tlogo)
				if len(tlogo) > 0:
					document.add_paragraph("", style='Cover Subtitle')
					# document.add_picture(tlogo)
					p = document.add_paragraph("", style='Cover Subtitle')
					r = p.add_run()
					r.add_picture(tlogo)
					document.add_paragraph("", style='Cover Subtitle')
			#
			# Execution Date range
			#
			execdr = ""
			if base.rs_setting_get_int("showstarttime"):
				iST = base.rs_setting_get_starttime()
				fSD = "{}".format(base.report_formatdate(iST))
				fST = "{}".format(base.report_formattime(iST))

				execdr = "{} {}".format(fSD, fST)

			if base.rs_setting_get_int("showendtime"):
				iET = base.rs_setting_get_endtime()
				fED = "{}".format(base.report_formatdate(iET))
				fET = "{}".format(base.report_formattime(iET))

				if not base.rs_setting_get_int("showstarttime"):
					execdr = "{} {}".format(fED, fET)
				else:
					if fSD == fED:
						execdr = "{} - {}".format(execdr, fET)
					else:
						execdr = "{} - {} {}".format(execdr, fED, fET)

			document.add_paragraph("", style='Cover Subtitle')
			document.add_paragraph(execdr, style='Cover Subtitle')
			document.add_paragraph("", style='Cover Subtitle')
		else:
			newsectionpct = 1 / (len(sections) + 1)
			sectionpct = newsectionpct * sectionpct
			base.debugmsg(5, "sectionpct:", sectionpct)
			self.docx_sections_addheading(id)

			stype = base.report_item_get_type(id)
			base.debugmsg(5, "stype:", stype)
			if stype == "contents":
				self.docx_sections_contents(id)
			if stype == "note":
				self.docx_sections_note(id)
			if stype == "graph":
				self.docx_sections_graph(id)
			if stype == "table":
				self.docx_sections_table(id)
			if stype == "errors":
				self.docx_sections_errors(id)

		self.cg_data["docx"]["progress"] += sectionpct
		self.display_message("Generating Word Report {}%".format(int(round(self.cg_data["docx"]["progress"] * 100, 0))))

		if len(sections) > 0:
			for sect in sections:
				self.docx_add_sections(sect, sectionpct)

	def docx_sections_addheading(self, id):
		base.debugmsg(5, "id:", id)
		document = self.cg_data["docx"]["document"]

		level = base.report_sect_level(id)
		base.debugmsg(5, "level:", level)
		number = base.report_sect_number(id)
		base.debugmsg(5, "number:", number)
		name = base.report_item_get_name(id)
		base.debugmsg(5, "name:", name)

		heading_text = "{} {}".format(number, name)

		base.debugmsg(5, "heading_text:", heading_text, "	level:", level)
		hdpg = document.add_heading(heading_text, level)
		stylename = "Heading {}".format(level)
		base.debugmsg(5, "stylename:", stylename)
		hdpg.style = stylename
		# document.add_paragraph("", style='Normal')

	def docx_sections_contents(self, id):
		base.debugmsg(5, "id:", id)

		mode = base.rt_contents_get_mode(id)
		level = base.rt_contents_get_level(id)

		base.debugmsg(5, "mode:", mode, "	level:", level)
		# fmode = None
		# if mode == "Table Of Contents":
		# 	fmode = None
		# if mode == "Table of Graphs":
		# 	fmode = "graph"
		# if mode == "Table Of Tables":
		# 	fmode = "table"

		document = self.cg_data["docx"]["document"]
		# document.add_paragraph("", style='Normal')

		paragraph = document.add_paragraph(style='Normal')
		run = paragraph.add_run()
		fldChar = OxmlElement('w:fldChar')  # creates a new element
		fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
		instrText = OxmlElement('w:instrText')
		instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
		if mode == "Table Of Contents":
			instrText.text = 'TOC \\o "1-' + str(level) + '" \\h \\z \\u'   # change 1-3 depending on heading levels you need

		# \c "SEQIdentifier"
		# Lists figures, tables, charts, or other items that are numbered by a SEQ (Sequence) field. Word uses SEQ fields to number items captioned with the Caption command (References > Insert Caption). SEQIdentifier, which corresponds to the caption label, must match the identifier in the SEQ field. For example, { TOC \c "tables" } lists all numbered tables.

		if mode == "Table of Graphs":
			instrText.text = 'TOC \\o "1-' + str(level) + '" \\c "figures"'

		if mode == "Table Of Tables":
			instrText.text = 'TOC \\o "1-' + str(level) + '" \\c "tables"'

		fldChar2 = OxmlElement('w:fldChar')
		fldChar2.set(qn('w:fldCharType'), 'separate')
		fldChar3 = OxmlElement('w:t')
		fldChar3.text = "Right-click to update field."
		fldChar2.append(fldChar3)

		fldChar4 = OxmlElement('w:fldChar')
		fldChar4.set(qn('w:fldCharType'), 'end')

		r_element = run._r
		r_element.append(fldChar)
		r_element.append(instrText)
		r_element.append(fldChar2)
		r_element.append(fldChar4)

	def docx_sections_note(self, id):
		base.debugmsg(5, "id:", id)

		document = self.cg_data["docx"]["document"]

		notebody = base.rt_note_get(id)
		notebody = notebody.replace("\r\n", "\n")
		notebody = notebody.replace("\r", "\n")
		notelist = notebody.split("\n")
		for line in notelist:
			document.add_paragraph(line, style='Normal')

	def docx_sections_graph(self, id):
		base.debugmsg(5, "id:", id)
		pid, idl, idr = base.rt_graph_LR_Ids(id)
		base.debugmsg(5, "pid:", pid, "	idl:", idl, "	idr:", idr)

		document = self.cg_data["docx"]["document"]
		# document.add_paragraph("", style='Normal')

		axisenl = base.rt_graph_get_axisen(idl)
		axisenr = base.rt_graph_get_axisen(idr)

		datatypel = base.rt_graph_get_dt(idl)
		datatyper = base.rt_graph_get_dt(idr)

		tz = zoneinfo.ZoneInfo(base.rs_setting_get_timezone())

		if datatypel == "SQL":
			sqll = base.rt_graph_get_sql(idl)
		else:
			sqll = base.rt_graph_generate_sql(idl)

		if datatyper == "SQL":
			sqlr = base.rt_graph_get_sql(idr)
		else:
			sqlr = base.rt_graph_generate_sql(idr)

		gphdpi = 72
		# gphdpi = 100
		fig = Figure(dpi=gphdpi)
		axisl = fig.add_subplot(1, 1, 1)
		axisl.grid(True, 'major', 'x')
		axisr = axisl.twinx()
		fig.autofmt_xdate(bottom=0.2, rotation=30, ha='right')

		canvas = FigureCanvas(fig)

		# https://stackoverflow.com/questions/57316491/how-to-convert-matplotlib-figure-to-pil-image-object-without-saving-image
		axisl.tick_params(labelleft=False, length=0)
		axisr.tick_params(labelright=False, length=0)

		try:
			canvas.draw()
		except Exception as e:
			base.debugmsg(5, "canvas.draw() Exception:", e)
		fig.set_tight_layout(True)

		dodraw = False
		graphdata = {}

		if (sqll is not None and len(sqll.strip()) > 0) or (sqlr is not None and len(sqlr.strip()) > 0):

			if sqll is not None and len(sqll.strip()) > 0 and axisenl > 0:
				base.debugmsg(7, "sqll:", sqll)
				key = "{}_{}".format(idl, base.report_item_get_changed(idl))
				base.dbqueue["Read"].append({"SQL": sqll, "KEY": key})
				while key not in base.dbqueue["ReadResult"]:
					time.sleep(0.1)

				gdata = base.dbqueue["ReadResult"][key]

				if datatypel == "Plan":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_plan(idl, gdata)
				if datatypel == "Monitoring":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_monitoring(idl, gdata)

				base.debugmsg(9, "gdata:", gdata)

				for row in gdata:
					base.debugmsg(9, "row:", row)
					if 'Name' in row:
						name = row['Name']
						base.debugmsg(9, "name:", name)
						if name not in graphdata:
							graphdata[name] = {}

							colour = base.named_colour(name)
							base.debugmsg(8, "name:", name, "	colour:", colour)
							graphdata[name]["Colour"] = colour
							graphdata[name]["Axis"] = "axisL"
							# self.contentdata[id]["graphdata"][name]["Time"] = []
							graphdata[name]["objTime"] = []
							graphdata[name]["Values"] = []

						graphdata[name]["objTime"].append(datetime.fromtimestamp(row["Time"], tz))
						graphdata[name]["Values"].append(base.rt_graph_floatval(row["Value"]))
					else:
						break

			if sqlr is not None and len(sqlr.strip()) > 0 and axisenr > 0:
				base.debugmsg(7, "sqlr:", sqlr)
				key = "{}_{}".format(idr, base.report_item_get_changed(idr))
				base.dbqueue["Read"].append({"SQL": sqlr, "KEY": key})
				while key not in base.dbqueue["ReadResult"]:
					time.sleep(0.1)

				gdata = base.dbqueue["ReadResult"][key]

				if datatyper == "Plan":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_plan(idr, gdata)
				if datatyper == "Monitoring":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_monitoring(idr, gdata)

				base.debugmsg(9, "gdata:", gdata)

				for row in gdata:
					base.debugmsg(9, "row:", row)
					if 'Name' in row:
						name = row['Name']
						base.debugmsg(9, "name:", name)
						if name not in graphdata:
							graphdata[name] = {}

							colour = base.named_colour(name)
							base.debugmsg(8, "name:", name, "	colour:", colour)
							graphdata[name]["Colour"] = colour
							graphdata[name]["Axis"] = "axisR"
							# self.contentdata[id]["graphdata"][name]["Time"] = []
							graphdata[name]["objTime"] = []
							graphdata[name]["Values"] = []

						graphdata[name]["objTime"].append(datetime.fromtimestamp(row["Time"], tz))
						graphdata[name]["Values"].append(base.rt_graph_floatval(row["Value"]))
					else:
						break

			base.debugmsg(9, "graphdata:", graphdata)

			for name in graphdata:
				base.debugmsg(7, "name:", name)

				axis = "axisL"
				if "Axis" in graphdata[name]:
					axis = graphdata[name]["Axis"]

				if len(graphdata[name]["Values"]) > 1 and len(graphdata[name]["Values"]) == len(graphdata[name]["objTime"]):
					try:
						if axis == "axisL":
							axisl.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name)
						elif axis == "axisR":
							axisr.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name)
						dodraw = True
					except Exception as e:
						base.debugmsg(7, "axis.plot() Exception:", e)

				if len(graphdata[name]["Values"]) == 1 and len(graphdata[name]["Values"]) == len(graphdata[name]["objTime"]):
					try:
						if axis == "axisL":
							axisl.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name, marker='o')
						elif axis == "axisR":
							axisr.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name, marker='o')
						dodraw = True
					except Exception as e:
						base.debugmsg(7, "axis.plot() Exception:", e)

			if dodraw:

				# Left axis Limits
				if axisenl > 0:
					axisl.grid(True, 'major', 'y')
					axisl.tick_params(labelleft=True, length=5)

					SMetric = "Other"
					if datatypel == "Metric":
						SMetric = base.rt_table_get_sm(idl)
					base.debugmsg(8, "SMetric:", SMetric)
					if SMetric in ["Load", "CPU", "MEM", "NET"]:
						axisl.set_ylim(0, 100)
					else:
						axisl.set_ylim(0)

				# Right axis Limits
				if axisenr > 0:
					axisr.grid(True, 'major', 'y')
					axisr.tick_params(labelright=True, length=5)

					SMetric = "Other"
					if datatyper == "Metric":
						SMetric = base.rt_table_get_sm(idr)
					base.debugmsg(8, "SMetric:", SMetric)
					if SMetric in ["Load", "CPU", "MEM", "NET"]:
						axisr.set_ylim(0, 100)
					else:
						axisr.set_ylim(0)

				fig.set_tight_layout(True)
				fig.autofmt_xdate(bottom=0.2, rotation=30, ha='right')
				try:
					canvas.draw()
				except Exception as e:
					base.debugmsg(5, "canvas.draw() Exception:", e)

				# works but messy createing files we then need to delete
				# filename = "{}.png".format(id)
				# fig.savefig(filename)
				# self.xhtml_sections_fileimg(body, id, filename)

				buf = BytesIO()
				fig.savefig(buf)
				buf.seek(0)
				pic = document.add_picture(buf)
				base.debugmsg(5, "pic:", pic)
				# base.debugmsg(5, "pic.parent:", pic.parent)
				document.paragraphs[-1].paragraph_format.left_indent = Cm(-0.5)

	def docx_sections_table(self, id):
		base.debugmsg(5, "id:", id)

		document = self.cg_data["docx"]["document"]
		# document.add_paragraph("", style='Normal')

		datatype = base.rt_table_get_dt(id)
		if datatype == "SQL":
			sql = base.rt_table_get_sql(id)
		else:
			sql = base.rt_table_generate_sql(id)
		colours = base.rt_table_get_colours(id)

		if sql is not None and len(sql.strip()) > 0:
			base.debugmsg(8, "sql:", sql)
			key = "{}_{}".format(id, base.report_item_get_changed(id))
			base.dbqueue["Read"].append({"SQL": sql, "KEY": key})
			while key not in base.dbqueue["ReadResult"]:
				time.sleep(0.1)

			tdata = base.dbqueue["ReadResult"][key]
			if datatype == "Plan":
				base.debugmsg(9, "tdata before:", tdata)
				tdata = base.table_postprocess_data_plan(id, tdata)
			if datatype == "Monitoring":
				base.debugmsg(9, "tdata before:", tdata)
				tdata = base.table_postprocess_data_monitoring(id, tdata)
			base.debugmsg(8, "tdata:", tdata)

			if len(tdata) > 0:
				# table headers
				cols = list(tdata[0].keys())
				base.debugmsg(7, "cols:", cols)

				# numcols = len(cols)
				numcols = 1
				cellcol = 0
				cellrow = 0
				if colours:
					# numcols += 1
					cellcol += 1
				# if "Colour" in cols:
				# 	numcols -= 1

				table = document.add_table(rows=1, cols=numcols)
				# Table Grid Light
				# Table Grid
				# table.style = "Table Grid Light"
				table.style = document.styles['Table Grid']

				if colours:
					table.columns[cellcol - 1].width = Cm(0.5)
					table.rows[cellrow].cells[cellcol - 1].paragraphs[0].style = "Table Header"

				cw = 5
				for col in cols:
					if col not in ["Colour"]:
						show = base.report_item_get_bool_def1(id, base.rt_table_ini_colname(f"{col} Show"))
						if show:
							if cellcol > 0:
								table.add_column(width=1)
							dispname = base.rt_table_get_colname(id, col)
							table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
							table.rows[cellrow].cells[cellcol].paragraphs[0].text = dispname.strip()
							table.columns[cellcol].width = Cm(cw)
							if cw > 2:
								cw = 1.7
							if cellcol > 5:
								cw = 1.1
							cellcol += 1

				# table rows
				for row in tdata:

					cellcol = 0
					cellrow += 1

					vals = list(row.values())
					base.debugmsg(7, "vals:", vals)
					table.add_row()
					if colours:

						base.debugmsg(9, "row:", row)
						if "Colour" in row:
							label = row["Colour"]
						else:
							label = row[cols[0]]
						base.debugmsg(9, "label:", label)
						colour = base.named_colour(label)
						base.debugmsg(9, "colour:", colour)

						swatch = Image.new(mode="RGB", size=(100, 100), color=colour)

						buffered = BytesIO()
						# swatch.save(buffered, format=swatch.format)
						swatch.save(buffered, format="PNG")
						buffered.seek(0)
						# img_byte = buffered.getvalue()
						#
						# srcdata += base64.b64encode(img_byte).decode()

						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						# table.rows[cellrow].cells[cellcol].paragraphs[0].paragraph_format.space_before = Pt(0)
						table.rows[cellrow].cells[cellcol].paragraphs[0].paragraph_format.left_indent = Cm(-0.50)

						run = table.rows[cellrow].cells[cellcol].paragraphs[0].add_run()
						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"

						run.add_picture(buffered, width=Cm(0.5))

						cellcol += 1

					# for val in vals:
					for col in cols:
						if col not in ["Colour"]:
							show = base.report_item_get_bool_def1(id, base.rt_table_ini_colname(f"{col} Show"))
							if show:
								val = str(row[col]).strip()
								val = base.illegal_xml_chars_re.sub('', val)
								base.debugmsg(8, "val:", val)

								# table.rows[cellrow].cells[cellcol].text = str(val)
								# table.rows[cellrow].cells[cellcol].add_paragraph(text=str(val), style="Table Cell")
								table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
								table.rows[cellrow].cells[cellcol].paragraphs[0].text = val

								tcw = int(table.columns[cellcol].width.cm) + 1
								# base.debugmsg(5, "tcw:", tcw)
								if tcw > 5:
									table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

								cellcol += 1

				table.autofit = True
				# table.style.paragraph_format.left_indent = Cm(-0.50)
				# table.style.paragraph_format.right_indent = Cm(-0.50)

	def docx_sections_errors(self, id):
		base.debugmsg(8, "id:", id)

		document = self.cg_data["docx"]["document"]

		imgsizew = 1400000 * 3
		cellcol = 0
		cellrow = -1

		table = document.add_table(rows=1, cols=7)
		table.style = document.styles['Table Grid']
		# table.alignment = WD_TABLE_ALIGNMENT.LEFT
		table.allow_autofit = True

		showimages = base.rt_errors_get_images(id)
		base.debugmsg(5, "showimages:", showimages)
		grouprn = base.rt_errors_get_group_rn(id)
		base.debugmsg(5, "grouprn:", grouprn)
		groupet = base.rt_errors_get_group_et(id)
		base.debugmsg(5, "groupet:", groupet)

		lbl_Result = base.rt_errors_get_label(id, "lbl_Result")
		lbl_Test = base.rt_errors_get_label(id, "lbl_Test")
		lbl_Script = base.rt_errors_get_label(id, "lbl_Script")
		lbl_Error = base.rt_errors_get_label(id, "lbl_Error")
		lbl_Count = base.rt_errors_get_label(id, "lbl_Count")
		lbl_Screenshot = base.rt_errors_get_label(id, "lbl_Screenshot")
		lbl_NoScreenshot = base.rt_errors_get_label(id, "lbl_NoScreenshot")

		pctalike = 0.80
		base.debugmsg(5, "pctalike:", pctalike)

		base.rt_errors_get_data(id)

		grpdata = {}
		if grouprn or groupet:
			grpdata = {}
			grpdata["resultnames"] = {}
			grpdata["errortexts"] = {}

			keys = list(base.reportdata[id].keys())
			for key in keys:
				base.debugmsg(5, "key:", key)
				if key != "key":
					rdata = base.reportdata[id][key]

					if grouprn:
						result_name = rdata['result_name']
						matches = difflib.get_close_matches(result_name, list(grpdata["resultnames"].keys()), cutoff=pctalike)
						base.debugmsg(5, "matches:", matches)
						if len(matches) > 0:
							result_name = matches[0]
							basekey = grpdata["resultnames"][result_name]["keys"][0]
							base.debugmsg(5, "basekey:", basekey)
							grpdata["resultnames"][result_name]["keys"].append(key)
						else:
							grpdata["resultnames"][result_name] = {}
							grpdata["resultnames"][result_name]["keys"] = []
							grpdata["resultnames"][result_name]["keys"].append(key)
							grpdata["resultnames"][result_name]["errortexts"] = {}

						if groupet and 'error' in rdata:
							errortext = rdata['error']
							# errortext_sub = errortext.split(r'\n')[0]
							errortext_sub = errortext.splitlines()[0]
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							matcheset = difflib.get_close_matches(errortext_sub, list(grpdata["resultnames"][result_name]["errortexts"].keys()), cutoff=pctalike)
							base.debugmsg(5, "matcheset:", matcheset)
							if len(matcheset) > 0:
								errortext = matcheset[0]
								errortext_sub = errortext.splitlines()[0]
								baseid = grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"][0]
								base.debugmsg(5, "baseid:", baseid)
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"].append(key)
							else:
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub] = {}
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"] = []
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"].append(key)

					if groupet and 'error' in rdata:
						errortext = rdata['error']
						errortext_sub = errortext.splitlines()[0]
						base.debugmsg(5, "errortext_sub:", errortext_sub)
						matches = difflib.get_close_matches(errortext_sub, list(grpdata["errortexts"].keys()), cutoff=pctalike)
						base.debugmsg(5, "matches:", matches)
						if len(matches) > 0:
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							errortext = matches[0]
							base.debugmsg(5, "errortext:", errortext)
							baseid = grpdata["errortexts"][errortext]["keys"][0]
							base.debugmsg(5, "baseid:", baseid)
							grpdata["errortexts"][errortext]["keys"].append(key)
						else:
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							grpdata["errortexts"][errortext_sub] = {}
							grpdata["errortexts"][errortext_sub]["keys"] = []
							grpdata["errortexts"][errortext_sub]["keys"].append(key)

			resultnames = grpdata["resultnames"]
			base.debugmsg(5, "resultnames:", resultnames)
			errortexts = grpdata["errortexts"]
			base.debugmsg(5, "errortexts:", errortexts)

		if grouprn:
			# add 2 columns for count
			table.add_column(Cm(1.8))
			table.add_column(Cm(0.8))

			for result_name in list(grpdata["resultnames"].keys()):
				basekey = grpdata["resultnames"][result_name]["keys"][0]
				base.debugmsg(5, "basekey:", basekey)
				rdata = base.reportdata[id][basekey]

				cellcol = 0
				cellrow += 1

				if cellrow > 0:
					table.add_row()

				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Result)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
				a = table.cell(cellrow, cellcol)
				b = table.cell(cellrow, cellcol + 1)
				a.merge(b)
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['result_name']
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
				table.rows[cellrow].cells[cellcol].paragraphs[0].FitText = True

				cellcol += 2
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Test)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['test_name']
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Script)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
				# table.columns[cellcol].width = Cm(1.8)

				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['script']
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Count)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
				# table.columns[cellcol].width = Cm(1.8)

				count = len(grpdata["resultnames"][result_name]["keys"])
				base.debugmsg(5, "count:", count)
				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = str(count)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				if groupet:
					for errortext in list(grpdata["resultnames"][result_name]["errortexts"].keys()):
						basekey = grpdata["resultnames"][result_name]["errortexts"][errortext]["keys"][0]
						base.debugmsg(5, "basekey:", basekey)
						rdata = base.reportdata[id][basekey]

						cellcol = 0
						cellrow += 1
						table.add_row()

						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Error)
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						cellcol += 1
						a = table.cell(cellrow, cellcol)
						b = table.cell(cellrow, cellcol + 5)
						a.merge(b)

						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
						if 'error' in rdata:
							table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['error']
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						cellcol += 6
						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Count)
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
						# table.columns[cellcol].width = Cm(1.8)

						count = len(grpdata["resultnames"][result_name]["errortexts"][errortext]["keys"])
						base.debugmsg(5, "count:", count)
						cellcol += 1
						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = str(count)
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						if showimages:
							cellcol = 0
							cellrow += 1
							table.add_row()

							table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
							table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Screenshot)
							table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

							cellcol += 1
							a = table.cell(cellrow, cellcol)
							b = table.cell(cellrow, cellcol + 7)
							a.merge(b)
							table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
							if 'image_file' in rdata:
								run = table.rows[cellrow].cells[cellcol].paragraphs[0].add_run()
								run.add_picture(rdata['image_file'], width=imgsizew)
							else:
								table.rows[cellrow].cells[cellcol].paragraphs[0].text = lbl_NoScreenshot
								table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				else:
					for keyi in grpdata["resultnames"][result_name]["keys"]:
						rdata = base.reportdata[id][keyi]

						cellcol = 0
						cellrow += 1
						table.add_row()

						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Error)
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						cellcol += 1
						a = table.cell(cellrow, cellcol)
						b = table.cell(cellrow, cellcol + 7)
						a.merge(b)

						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
						if 'error' in rdata:
							table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['error']
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						if showimages:
							cellcol = 0
							cellrow += 1
							table.add_row()

							table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
							table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Screenshot)
							table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

							cellcol += 1
							a = table.cell(cellrow, cellcol)
							b = table.cell(cellrow, cellcol + 7)
							a.merge(b)
							table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
							if 'image_file' in rdata:
								run = table.rows[cellrow].cells[cellcol].paragraphs[0].add_run()
								run.add_picture(rdata['image_file'], width=imgsizew)
							else:
								table.rows[cellrow].cells[cellcol].paragraphs[0].text = lbl_NoScreenshot
								table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

		if groupet and not grouprn:
			# add 2 columns for count
			table.add_column(Cm(1.8))
			table.add_column(Cm(0.8))

			for errortext in list(grpdata["errortexts"].keys()):
				basekey = grpdata["errortexts"][errortext]["keys"][0]
				base.debugmsg(5, "basekey:", basekey)
				rdata = base.reportdata[id][basekey]

				cellcol = 0
				cellrow += 1

				if cellrow > 0:
					table.add_row()

				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Error)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				cellcol += 1
				a = table.cell(cellrow, cellcol)
				b = table.cell(cellrow, cellcol + 5)
				a.merge(b)

				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
				if 'error' in rdata:
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['error']
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				cellcol += 6
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Count)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				count = len(grpdata["errortexts"][errortext]["keys"])
				base.debugmsg(5, "count:", count)
				cellcol += 1
				table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
				table.rows[cellrow].cells[cellcol].paragraphs[0].text = str(count)
				table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

				if showimages:
					cellcol = 0
					cellrow += 1
					table.add_row()

					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Screenshot)
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

					cellcol += 1
					a = table.cell(cellrow, cellcol)
					b = table.cell(cellrow, cellcol + 7)
					a.merge(b)
					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
					if 'image_file' in rdata:
						run = table.rows[cellrow].cells[cellcol].paragraphs[0].add_run()
						run.add_picture(rdata['image_file'], width=imgsizew)
					else:
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = lbl_NoScreenshot
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

		if not grouprn and not groupet:
			keys = list(base.reportdata[id].keys())
			for key in keys:
				base.debugmsg(5, "key:", key)
				if key != "key":
					rdata = base.reportdata[id][key]

					cellcol = 0
					cellrow += 1

					if cellrow > 0:
						table.add_row()

					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Result)
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
					# table.columns[cellcol].width = Cm(1.8)

					cellcol += 1
					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
					a = table.cell(cellrow, cellcol)
					b = table.cell(cellrow, cellcol + 1)
					a.merge(b)
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['result_name']
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
					table.rows[cellrow].cells[cellcol].paragraphs[0].FitText = True

					cellcol += 2
					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Test)
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
					# table.columns[cellcol].width = Cm(1.8)

					cellcol += 1
					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['test_name']
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

					cellcol += 1
					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Script)
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
					# table.columns[cellcol].width = Cm(1.8)

					cellcol += 1
					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['script']
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

					cellcol = 0
					cellrow += 1
					table.add_row()

					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
					table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Error)
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

					cellcol += 1
					a = table.cell(cellrow, cellcol)
					b = table.cell(cellrow, cellcol + 5)
					a.merge(b)

					table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
					if 'error' in rdata:
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = rdata['error']
					table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

					if showimages:
						cellcol = 0
						cellrow += 1
						table.add_row()

						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Header"
						table.rows[cellrow].cells[cellcol].paragraphs[0].text = "{}:".format(lbl_Screenshot)
						table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

						cellcol += 1
						a = table.cell(cellrow, cellcol)
						b = table.cell(cellrow, cellcol + 5)
						a.merge(b)
						table.rows[cellrow].cells[cellcol].paragraphs[0].style = "Table Cell"
						if 'image_file' in rdata:
							run = table.rows[cellrow].cells[cellcol].paragraphs[0].add_run()
							run.add_picture(rdata['image_file'], width=imgsizew)
						else:
							table.rows[cellrow].cells[cellcol].paragraphs[0].text = lbl_NoScreenshot
							table.rows[cellrow].cells[cellcol].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

	#
	# 	MS Excel
	#
	# https://xlsxwriter.readthedocs.io/introduction.html
	# https://openpyxl.readthedocs.io/en/stable/index.html

	def xlsx_configure_style(self):

		basecolour = '000000'
		highlightcolour = base.rs_setting_get_hcolour().replace("#", "")
		fontname = base.rs_setting_get_font()
		fontsize = base.rs_setting_get_fontsize()

		wb = self.cg_data["xlsx"]["Workbook"]

		default = openpyxl.styles.NamedStyle(name="Default")
		default.font.name = fontname
		default.font.size = fontsize
		default.font.color = basecolour
		wb.add_named_style(default)

		# highlight = openpyxl.styles.NamedStyle(name="Highlight")
		highlight = copy(default)
		highlight.name = "Highlight"
		# highlight.font.name = fontname
		# highlight.font.size = fontsize
		highlight.font.color = highlightcolour
		wb.add_named_style(highlight)

		# title = openpyxl.styles.NamedStyle(name="CoverTitle")
		title = copy(default)
		title.name = "CoverTitle"
		# title.font.name = fontname
		title.font.size = fontsize * 2
		# title.font.color = basecolour
		title.alignment.horizontal = 'center'
		title.alignment.wrapText = True
		wb.add_named_style(title)
		base.debugmsg(9, "title:", title.name, title.font.name, title.font.size)

		# subtitle = openpyxl.styles.NamedStyle(name="CoverSubTitle")
		subtitle = copy(title)
		subtitle.name = "CoverSubTitle"
		# subtitle.font.name = fontname
		subtitle.font.size = fontsize * 1.5
		# subtitle.font.color = basecolour
		# subtitle.alignment.horizontal = 'center'
		# subtitle.alignment.wrapText = True
		wb.add_named_style(subtitle)
		base.debugmsg(9, "subtitle:", subtitle.name, subtitle.font.name, subtitle.font.size)
		base.debugmsg(9, "title:", title.name, title.font.name, title.font.size)
		base.debugmsg(9, "highlight:", highlight.name, highlight.font.name, highlight.font.size)
		base.debugmsg(9, "default:", default.name, default.font.name, default.font.size)

		headings = {}
		fm = 2
		for i in range(6):
			base.debugmsg(7, "i:", i, i + 1)
			hnum = i + 1
			headings[hnum] = copy(highlight)
			headings[hnum].name = "Heading " + str(hnum)
			headings[hnum].font.size = int(fontsize * fm)
			wb.add_named_style(headings[hnum])

			fm -= 0.2

		# Table Heading
		side = openpyxl.styles.borders.Side(style="tblside", color='CCCCCC', border_style='thin')  # 'thin' 'hair' 'medium'
		borders = openpyxl.styles.borders.Border(left=side, right=side, top=side, bottom=side)
		tableh = copy(highlight)
		tableh.name = "Table Heading"
		tableh.border = borders
		tableh.alignment.vertical = "top"
		wb.add_named_style(tableh)

		# Table Data
		tabled = copy(default)
		tabled.name = "Table Data"
		tabled.border = borders
		tabled.alignment.vertical = "top"
		wb.add_named_style(tabled)

	def xlsx_add_sections(self, id, sectionpct):
		base.debugmsg(7, "id:", id, "	sectionpct:", sectionpct)

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active

		base.debugmsg(9, "ws:", ws)
		base.debugmsg(9, "ws.title:", ws.title)

		sections = base.report_get_order(id)
		base.debugmsg(9, "sections:", sections)

		if id == "TOP":

			ws.title = "Cover"

			rownum = 0

			#
			# Title
			#
			titletxt = base.rs_setting_get_title()
			rownum = 3
			colspan = 9

			ws.merge_cells(start_row=rownum, start_column=1, end_row=rownum, end_column=colspan)
			titlecell = ws.cell(column=1, row=rownum, value=titletxt)
			titlecell.style = "CoverTitle"

			fontsize = base.rs_setting_get_fontsize()
			rd = ws.row_dimensions[rownum]
			rd.height = fontsize * 5

			#
			# Logo
			#
			rownum = 5
			base.debugmsg(7, "showtlogo:", base.rs_setting_get_int("showtlogo"))
			if base.rs_setting_get_int("showtlogo"):

				tlogo = base.rs_setting_get_file("tlogo")
				base.debugmsg(7, "tlogo:", tlogo)
				if len(tlogo) > 0:
					img = openpyxl.drawing.image.Image(tlogo)
					cellname = ws.cell(row=rownum, column=1).coordinate
					ws.add_image(img, cellname)

			#
			# Execution Date range
			#
			rownum = 20

			execdr = ""
			if base.rs_setting_get_int("showstarttime"):
				iST = base.rs_setting_get_starttime()
				fSD = "{}".format(base.report_formatdate(iST))
				fST = "{}".format(base.report_formattime(iST))

				execdr = "{} {}".format(fSD, fST)

			if base.rs_setting_get_int("showendtime"):
				iET = base.rs_setting_get_endtime()
				fED = "{}".format(base.report_formatdate(iET))
				fET = "{}".format(base.report_formattime(iET))

				if not base.rs_setting_get_int("showstarttime"):
					execdr = "{} {}".format(fED, fET)
				else:
					if fSD == fED:
						execdr = "{} - {}".format(execdr, fET)
					else:
						execdr = "{} - {} {}".format(execdr, fED, fET)

			ws.merge_cells(start_row=rownum, start_column=1, end_row=rownum, end_column=colspan)
			subtitlecell = ws.cell(column=1, row=rownum, value=execdr)
			subtitlecell.style = "CoverSubTitle"

			rd = ws.row_dimensions[rownum]
			rd.height = fontsize * 2

		else:
			newsectionpct = 1 / (len(sections) + 1)
			sectionpct = newsectionpct * sectionpct
			base.debugmsg(8, "sectionpct:", sectionpct)
			self.xlsx_sections_addheading(id)

			stype = base.report_item_get_type(id)
			base.debugmsg(7, "stype:", stype)
			if stype == "contents":
				self.xlsx_sections_contents(id)
				pass
			if stype == "note":
				self.xlsx_sections_note(id)
				pass
			if stype == "graph":
				self.xlsx_sections_graph(id)
				pass
			if stype == "table":
				self.xlsx_sections_table(id)
				pass
			if stype == "errors":
				self.xlsx_sections_errors(id)
				pass

		self.cg_data["xlsx"]["progress"] += sectionpct
		self.display_message("Generating Excel Report {}%".format(int(round(self.cg_data["xlsx"]["progress"] * 100, 0))))

		if len(sections) > 0:
			for sect in sections:
				self.xlsx_add_sections(sect, sectionpct)

	def xlsx_sections_addheading(self, id):
		base.debugmsg(8, "id:", id)

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active

		base.debugmsg(9, "ws:", ws)
		base.debugmsg(9, "ws.title:", ws.title)

		base.debugmsg(9, "ws.active_cell:", ws.active_cell, "	ws.selected_cell:", ws.selected_cell)

		# acell = ws.cell(ws.active_cell)
		acell = ws[ws.active_cell]
		rownum = acell.row

		level = base.report_sect_level(id)
		base.debugmsg(8, "level:", level)
		number = base.report_sect_number(id)
		base.debugmsg(9, "number:", number)
		name = base.report_item_get_name(id)
		base.debugmsg(9, "name:", name)

		heading_text = "{} {}".format(number, name)

		if level == 1:
			ws = wb.create_sheet(title=heading_text)
			for wsi in wb.worksheets:
				if wsi.title == heading_text:
					wb.active = wsi

			rownum = 1

		titlecell = ws.cell(column=1, row=rownum, value=heading_text)
		titlecell.style = "Heading " + str(level)

		if level < 5:
			fontsize = base.rs_setting_get_fontsize()
			rd = ws.row_dimensions[rownum]
			rd.height = fontsize * 2

		# base.debugmsg(5, "ws.active_cell:", ws.active_cell, "	ws.selected_cell:", ws.selected_cell)

		rownum += 1
		self.xlsx_select_cell(1, rownum)

		# base.debugmsg(5, "ws.active_cell:", ws.active_cell, "	ws.selected_cell:", ws.selected_cell)

	def xlsx_select_cell(self, col, row):

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active
		nextcell = ws.cell(column=col, row=row)

		ws.sheet_view.selection[0].activeCell = nextcell.coordinate
		ws.sheet_view.selection[0].sqref = nextcell.coordinate

	def xlsx_sections_contents(self, id):
		base.debugmsg(8, "id:", id)

		mode = base.rt_contents_get_mode(id)
		maxlevel = base.rt_contents_get_level(id)
		fmode = None
		if mode == "Table Of Contents":
			fmode = None
		if mode == "Table of Graphs":
			fmode = "graph"
		if mode == "Table Of Tables":
			fmode = "table"

		self.xlsx_sections_contents_row("TOP", maxlevel, fmode)

	def xlsx_sections_contents_row(self, id, maxlevel, mode):

		display = True

		sections = base.report_get_order(id)
		base.debugmsg(9, "sections:", sections)
		level = base.report_sect_level(id)
		base.debugmsg(8, "level:", level)

		if id == "TOP":
			display = False

		if mode is not None:
			type = base.report_item_get_type(id)
			base.debugmsg(8, "type:", type)
			if mode != type:
				display = False

		if display:
			wb = self.cg_data["xlsx"]["Workbook"]
			ws = wb.active
			rownum = ws[ws.active_cell].row

			number = base.report_sect_number(id)
			base.debugmsg(9, "number:", number)
			name = base.report_item_get_name(id)
			base.debugmsg(9, "name:", name)
			type = base.report_item_get_type(id)
			base.debugmsg(9, "type:", type)

			heading_text = "{} {}".format(number, name)

			if level > 1:
				parentid = base.report_item_parent(id)
				parentlvl = base.report_sect_level(parentid)
				while parentlvl > 1:
					base.debugmsg(9, "parentid:", parentid)
					base.debugmsg(9, "parentlvl:", parentlvl)

					parentid = base.report_item_parent(parentid)
					parentlvl = base.report_sect_level(parentid)

				base.debugmsg(9, "parentid:", parentid)
				base.debugmsg(9, "parentlvl:", parentlvl)

				pnumber = base.report_sect_number(parentid)
				base.debugmsg(9, "pnumber:", pnumber)
				pname = base.report_item_get_name(parentid)
				base.debugmsg(9, "pname:", pname)

				parent_text = "{} {}".format(pnumber, pname)
			else:
				parent_text = heading_text

			rownum += 1
			self.xlsx_select_cell(1, rownum)
			c = ws[ws.active_cell]
			# =HYPERLINK(CONCAT("#'10 agents'!A",MATCH("10.8 selenium versions",'10 agents' A:A,0)),'10.8 selenium versions')
			match = "MATCH(\"" + heading_text + "\",'" + parent_text + "'!A:A,0)"
			base.debugmsg(9, "match:", match)
			concat = "CONCATENATE(\"#'" + parent_text + "'!A\"," + match + ")"
			base.debugmsg(9, "concat:", concat)
			hyper = "=HYPERLINK(" + concat + ",\"" + heading_text + "\")"
			base.debugmsg(8, "hyper:", hyper)
			c.style = "Default"
			c.value = hyper

		if level < maxlevel:
			if len(sections) > 0:
				for sect in sections:
					self.xlsx_sections_contents_row(sect, maxlevel, mode)

	def xlsx_sections_note(self, id):
		base.debugmsg(8, "id:", id)

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active
		rownum = ws[ws.active_cell].row

		rownum += 1
		self.xlsx_select_cell(1, rownum)

		notebody = base.rt_note_get(id)
		notebody = notebody.replace("\r\n", "\n")
		notebody = notebody.replace("\r", "\n")
		notelist = notebody.split("\n")
		for line in notelist:
			linecell = ws.cell(column=1, row=rownum, value=line)
			linecell.style = "Default"
			rownum += 1
			self.xlsx_select_cell(1, rownum)

	def xlsx_sections_graph(self, id):
		base.debugmsg(5, "id:", id)
		pid, idl, idr = base.rt_graph_LR_Ids(id)
		base.debugmsg(5, "pid:", pid, "	idl:", idl, "	idr:", idr)

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active
		rownum = ws[ws.active_cell].row

		rownum += 1
		self.xlsx_select_cell(1, rownum)

		axisenl = base.rt_graph_get_axisen(idl)
		axisenr = base.rt_graph_get_axisen(idr)

		datatypel = base.rt_graph_get_dt(idl)
		datatyper = base.rt_graph_get_dt(idr)

		tz = zoneinfo.ZoneInfo(base.rs_setting_get_timezone())

		if datatypel == "SQL":
			sqll = base.rt_graph_get_sql(idl)
		else:
			sqll = base.rt_graph_generate_sql(idl)

		if datatyper == "SQL":
			sqlr = base.rt_graph_get_sql(idr)
		else:
			sqlr = base.rt_graph_generate_sql(idr)

		gphdpi = 72
		# gphdpi = 100
		fig = Figure(dpi=gphdpi)
		axisl = fig.add_subplot(1, 1, 1)
		axisl.grid(True, 'major', 'x')
		axisr = axisl.twinx()
		fig.autofmt_xdate(bottom=0.2, rotation=30, ha='right')

		canvas = FigureCanvas(fig)

		# https://stackoverflow.com/questions/57316491/how-to-convert-matplotlib-figure-to-pil-image-object-without-saving-image
		axisl.tick_params(labelleft=False, length=0)
		axisr.tick_params(labelright=False, length=0)

		try:
			canvas.draw()
		except Exception as e:
			base.debugmsg(5, "canvas.draw() Exception:", e)
		fig.set_tight_layout(True)

		dodraw = False
		graphdata = {}

		if (sqll is not None and len(sqll.strip()) > 0) or (sqlr is not None and len(sqlr.strip()) > 0):

			if sqll is not None and len(sqll.strip()) > 0 and axisenl > 0:
				base.debugmsg(7, "sqll:", sqll)
				key = "{}_{}".format(idl, base.report_item_get_changed(idl))
				base.dbqueue["Read"].append({"SQL": sqll, "KEY": key})
				while key not in base.dbqueue["ReadResult"]:
					time.sleep(0.1)

				gdata = base.dbqueue["ReadResult"][key]

				if datatypel == "Plan":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_plan(idl, gdata)
				if datatypel == "Monitoring":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_monitoring(idl, gdata)

				base.debugmsg(9, "gdata:", gdata)

				for row in gdata:
					base.debugmsg(9, "row:", row)
					if 'Name' in row:
						name = row['Name']
						base.debugmsg(9, "name:", name)
						if name not in graphdata:
							graphdata[name] = {}

							colour = base.named_colour(name)
							base.debugmsg(8, "name:", name, "	colour:", colour)
							graphdata[name]["Colour"] = colour
							graphdata[name]["Axis"] = "axisL"
							# self.contentdata[id]["graphdata"][name]["Time"] = []
							graphdata[name]["objTime"] = []
							graphdata[name]["Values"] = []

						graphdata[name]["objTime"].append(datetime.fromtimestamp(row["Time"], tz))
						graphdata[name]["Values"].append(base.rt_graph_floatval(row["Value"]))
					else:
						break

			if sqlr is not None and len(sqlr.strip()) > 0 and axisenr > 0:
				base.debugmsg(7, "sqlr:", sqlr)
				key = "{}_{}".format(idr, base.report_item_get_changed(idr))
				base.dbqueue["Read"].append({"SQL": sqlr, "KEY": key})
				while key not in base.dbqueue["ReadResult"]:
					time.sleep(0.1)

				gdata = base.dbqueue["ReadResult"][key]

				if datatyper == "Plan":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_plan(idr, gdata)
				if datatyper == "Monitoring":
					base.debugmsg(9, "gdata before:", gdata)
					gdata = base.graph_postprocess_data_monitoring(idr, gdata)

				base.debugmsg(9, "gdata:", gdata)

				for row in gdata:
					base.debugmsg(9, "row:", row)
					if 'Name' in row:
						name = row['Name']
						base.debugmsg(9, "name:", name)
						if name not in graphdata:
							graphdata[name] = {}

							colour = base.named_colour(name)
							base.debugmsg(8, "name:", name, "	colour:", colour)
							graphdata[name]["Colour"] = colour
							graphdata[name]["Axis"] = "axisR"
							# self.contentdata[id]["graphdata"][name]["Time"] = []
							graphdata[name]["objTime"] = []
							graphdata[name]["Values"] = []

						graphdata[name]["objTime"].append(datetime.fromtimestamp(row["Time"], tz))
						graphdata[name]["Values"].append(base.rt_graph_floatval(row["Value"]))
					else:
						break

			base.debugmsg(9, "graphdata:", graphdata)

			for name in graphdata:
				base.debugmsg(7, "name:", name)

				axis = "axisL"
				if "Axis" in graphdata[name]:
					axis = graphdata[name]["Axis"]

				if len(graphdata[name]["Values"]) > 1 and len(graphdata[name]["Values"]) == len(graphdata[name]["objTime"]):
					try:
						if axis == "axisL":
							axisl.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name)
						elif axis == "axisR":
							axisr.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name)
						dodraw = True
					except Exception as e:
						base.debugmsg(7, "axis.plot() Exception:", e)

				if len(graphdata[name]["Values"]) == 1 and len(graphdata[name]["Values"]) == len(graphdata[name]["objTime"]):
					try:
						if axis == "axisL":
							axisl.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name, marker='o')
						elif axis == "axisR":
							axisr.plot(graphdata[name]["objTime"], graphdata[name]["Values"], graphdata[name]["Colour"], label=name, marker='o')
						dodraw = True
					except Exception as e:
						base.debugmsg(7, "axis.plot() Exception:", e)

			if dodraw:

				# Left axis Limits
				if axisenl > 0:
					axisl.grid(True, 'major', 'y')
					axisl.tick_params(labelleft=True, length=5)

					SMetric = "Other"
					if datatypel == "Metric":
						SMetric = base.rt_table_get_sm(idl)
					base.debugmsg(8, "SMetric:", SMetric)
					if SMetric in ["Load", "CPU", "MEM", "NET"]:
						axisl.set_ylim(0, 100)
					else:
						axisl.set_ylim(0)

				# Right axis Limits
				if axisenr > 0:
					axisr.grid(True, 'major', 'y')
					axisr.tick_params(labelright=True, length=5)

					SMetric = "Other"
					if datatyper == "Metric":
						SMetric = base.rt_table_get_sm(idr)
					base.debugmsg(8, "SMetric:", SMetric)
					if SMetric in ["Load", "CPU", "MEM", "NET"]:
						axisr.set_ylim(0, 100)
					else:
						axisr.set_ylim(0)

				fig.set_tight_layout(True)
				fig.autofmt_xdate(bottom=0.2, rotation=30, ha='right')
				try:
					canvas.draw()
				except Exception as e:
					base.debugmsg(5, "canvas.draw() Exception:", e)

				buf = BytesIO()
				fig.savefig(buf)
				buf.seek(0)

				# rownum += 1
				img = openpyxl.drawing.image.Image(buf)
				cellname = ws.cell(row=rownum, column=2).coordinate
				ws.add_image(img, cellname)

				rownum += 19
				self.xlsx_select_cell(1, rownum)

	def xlsx_sections_table(self, id):
		base.debugmsg(8, "id:", id)

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active
		rownum = ws[ws.active_cell].row

		datatype = base.rt_table_get_dt(id)
		if datatype == "SQL":
			sql = base.rt_table_get_sql(id)
		else:
			sql = base.rt_table_generate_sql(id)
		colours = base.rt_table_get_colours(id)

		if sql is not None and len(sql.strip()) > 0:
			base.debugmsg(8, "sql:", sql)
			key = "{}_{}".format(id, base.report_item_get_changed(id))
			base.dbqueue["Read"].append({"SQL": sql, "KEY": key})
			while key not in base.dbqueue["ReadResult"]:
				time.sleep(0.1)

			tdata = base.dbqueue["ReadResult"][key]
			if datatype == "Plan":
				base.debugmsg(9, "tdata before:", tdata)
				tdata = base.table_postprocess_data_plan(id, tdata)
			if datatype == "Monitoring":
				base.debugmsg(9, "tdata before:", tdata)
				tdata = base.table_postprocess_data_monitoring(id, tdata)
			base.debugmsg(8, "tdata:", tdata)

			if len(tdata) > 0:
				# table headers
				cols = list(tdata[0].keys())
				base.debugmsg(7, "cols:", cols)

				rownum += 1
				numcols = len(cols)
				cellcol = 1
				if colours:
					# set first column narrow for colour swatch
					# ws.columns[1].width = 10
					ws.column_dimensions["A"].width = 3

					numcols += 1
					cellcol += 1

				else:
					ws.column_dimensions["A"].width = 3
					numcols += 1
					cellcol += 1

				for col in cols:
					if col not in ["Colour"]:
						show = base.report_item_get_bool_def1(id, base.rt_table_ini_colname(f"{col} Show"))
						if show:
							dispname = base.rt_table_get_colname(id, col)

							base.debugmsg(8, "col:", col, "	cellcol:", cellcol, "	rownum:", rownum)
							hcell = ws.cell(column=cellcol, row=rownum, value=dispname.strip())
							hcell.style = "Table Heading"

							neww = len(str(col.strip())) * 1.3
							base.debugmsg(9, "neww:", neww)
							ws.column_dimensions[hcell.column_letter].width = neww

							cellcol += 1

				# table rows
				for row in tdata:

					cellcol = 1
					rownum += 1

					vals = list(row.values())
					base.debugmsg(7, "vals:", vals)

					if colours:

						base.debugmsg(9, "row:", row)
						if "Colour" in row:
							label = row["Colour"]
						else:
							label = row[cols[0]]
						base.debugmsg(9, "label:", label)
						colour = base.named_colour(label).replace("#", "")
						base.debugmsg(9, "colour:", colour)
						dcell = ws.cell(column=cellcol, row=rownum)
						dcell.style = "Table Data"
						dcell.fill = openpyxl.styles.PatternFill("solid", fgColor=colour)

						cellcol += 1
					else:
						cellcol += 1

					# for val in vals:
					for col in cols:
						if col not in ["Colour"]:
							show = base.report_item_get_bool_def1(id, base.rt_table_ini_colname(f"{col} Show"))
							if show:
								val = str(row[col]).strip()
								val = base.illegal_xml_chars_re.sub('', val)
								base.debugmsg(8, "val:", val)

								dcell = ws.cell(column=cellcol, row=rownum, value=val)
								dcell.style = "Table Data"

								currw = ws.column_dimensions[dcell.column_letter].width
								base.debugmsg(9, "currw:", currw, "	len(val):", len(str(val)))
								neww = max(currw, len(str(val)))
								base.debugmsg(8, "neww:", neww)
								ws.column_dimensions[dcell.column_letter].width = neww

								cellcol += 1

		rownum += 2
		self.xlsx_select_cell(1, rownum)

	def xlsx_sections_errors(self, id):
		base.debugmsg(5, "id:", id)

		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active
		rownum = ws[ws.active_cell].row
		base.debugmsg(5, "rownum:", rownum)

		cellcol = 1
		base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
		ws.column_dimensions["A"].width = 3

		cellcol += 1
		base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)

		showimages = base.rt_errors_get_images(id)
		base.debugmsg(5, "showimages:", showimages)
		grouprn = base.rt_errors_get_group_rn(id)
		base.debugmsg(5, "grouprn:", grouprn)
		groupet = base.rt_errors_get_group_et(id)
		base.debugmsg(5, "groupet:", groupet)

		lbl_Result = base.rt_errors_get_label(id, "lbl_Result")
		lbl_Test = base.rt_errors_get_label(id, "lbl_Test")
		lbl_Script = base.rt_errors_get_label(id, "lbl_Script")
		lbl_Error = base.rt_errors_get_label(id, "lbl_Error")
		lbl_Count = base.rt_errors_get_label(id, "lbl_Count")
		lbl_Screenshot = base.rt_errors_get_label(id, "lbl_Screenshot")
		lbl_NoScreenshot = base.rt_errors_get_label(id, "lbl_NoScreenshot")

		pctalike = 0.80
		base.debugmsg(5, "pctalike:", pctalike)

		base.rt_errors_get_data(id)

		grpdata = {}
		if grouprn or groupet:
			grpdata = {}
			grpdata["resultnames"] = {}
			grpdata["errortexts"] = {}

			keys = list(base.reportdata[id].keys())
			for key in keys:
				base.debugmsg(5, "key:", key)
				if key != "key":
					rdata = base.reportdata[id][key]

					if grouprn:
						result_name = rdata['result_name']
						matches = difflib.get_close_matches(result_name, list(grpdata["resultnames"].keys()), cutoff=pctalike)
						base.debugmsg(5, "matches:", matches)
						if len(matches) > 0:
							result_name = matches[0]
							basekey = grpdata["resultnames"][result_name]["keys"][0]
							base.debugmsg(5, "basekey:", basekey)
							grpdata["resultnames"][result_name]["keys"].append(key)
						else:
							grpdata["resultnames"][result_name] = {}
							grpdata["resultnames"][result_name]["keys"] = []
							grpdata["resultnames"][result_name]["keys"].append(key)
							grpdata["resultnames"][result_name]["errortexts"] = {}

						if groupet and 'error' in rdata:
							errortext = rdata['error']
							# errortext_sub = errortext.split(r'\n')[0]
							errortext_sub = errortext.splitlines()[0]
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							matcheset = difflib.get_close_matches(errortext_sub, list(grpdata["resultnames"][result_name]["errortexts"].keys()), cutoff=pctalike)
							base.debugmsg(5, "matcheset:", matcheset)
							if len(matcheset) > 0:
								errortext = matcheset[0]
								errortext_sub = errortext.splitlines()[0]
								baseid = grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"][0]
								base.debugmsg(5, "baseid:", baseid)
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"].append(key)
							else:
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub] = {}
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"] = []
								grpdata["resultnames"][result_name]["errortexts"][errortext_sub]["keys"].append(key)

					if groupet and 'error' in rdata:
						errortext = rdata['error']
						errortext_sub = errortext.splitlines()[0]
						base.debugmsg(5, "errortext_sub:", errortext_sub)
						matches = difflib.get_close_matches(errortext_sub, list(grpdata["errortexts"].keys()), cutoff=pctalike)
						base.debugmsg(5, "matches:", matches)
						if len(matches) > 0:
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							errortext = matches[0]
							base.debugmsg(5, "errortext:", errortext)
							baseid = grpdata["errortexts"][errortext]["keys"][0]
							base.debugmsg(5, "baseid:", baseid)
							grpdata["errortexts"][errortext]["keys"].append(key)
						else:
							base.debugmsg(5, "errortext_sub:", errortext_sub)
							grpdata["errortexts"][errortext_sub] = {}
							grpdata["errortexts"][errortext_sub]["keys"] = []
							grpdata["errortexts"][errortext_sub]["keys"].append(key)

			resultnames = grpdata["resultnames"]
			base.debugmsg(5, "resultnames:", resultnames)
			errortexts = grpdata["errortexts"]
			base.debugmsg(5, "errortexts:", errortexts)

		if grouprn:

			for result_name in list(grpdata["resultnames"].keys()):
				basekey = grpdata["resultnames"][result_name]["keys"][0]
				base.debugmsg(5, "basekey:", basekey)
				rdata = base.reportdata[id][basekey]

				rownum += 1
				cellcol = 0

				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Result), "Table Heading", 0)
				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['result_name'], "Table Data", 0)

				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Test), "Table Heading", 0)
				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['test_name'], "Table Data", 0)

				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Script), "Table Heading", 0)
				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['script'], "Table Data", 0)

				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Count), "Table Heading", 0)
				count = len(grpdata["resultnames"][result_name]["keys"])
				cellcol += 1
				base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum)
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, str(count), "Table Data", 0)

				if groupet:
					for errortext in list(grpdata["resultnames"][result_name]["errortexts"].keys()):
						basekey = grpdata["resultnames"][result_name]["errortexts"][errortext]["keys"][0]
						base.debugmsg(5, "basekey:", basekey)
						rdata = base.reportdata[id][basekey]

						cellcol = 1
						rownum += 1
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Error), "Table Heading", 0)
						cellcol += 1
						if 'error' in rdata:
							self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['error'], "Table Data", 4)

						cellcol += 5
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Count), "Table Heading", 0)
						count = len(grpdata["resultnames"][result_name]["errortexts"][errortext]["keys"])
						cellcol += 1
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, str(count), "Table Data", 0)

						if showimages:
							cellcol = 1
							rownum += 1
							self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Screenshot), "Table Heading", 0)

							cellcol += 1

							if 'image_file' in rdata:
								self.xlsx_sections_errors_fill_cell(cellcol, rownum, " ", "Table Data", 4)
								img = openpyxl.drawing.image.Image(rdata['image_file'])
								cellname = ws.cell(row=rownum, column=cellcol).coordinate

								base.debugmsg(5, "img.width:", img.width, "	img.height:", img.height)
								# 								31.75					32.60
								# 								22.23					22.82	==> 70%
								newiw = 850
								ratio = newiw / img.width
								base.debugmsg(5, "ratio:", ratio)
								newih = img.height * ratio
								base.debugmsg(5, "newih:", newih)
								img.width = newiw
								img.height = newih

								newh = newih * 0.76
								# 43.44 cm ==> 32.95 cm	==>	76%
								base.debugmsg(5, "newh:", newh)
								ws.row_dimensions[rownum].height = newh

								ws.add_image(img, cellname)
							else:
								self.xlsx_sections_errors_fill_cell(cellcol, rownum, lbl_NoScreenshot, "Table Data", 4)

				else:
					for keyi in grpdata["resultnames"][result_name]["keys"]:
						rdata = base.reportdata[id][keyi]

						cellcol = 1
						rownum += 1

						self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Error), "Table Heading", 0)
						cellcol += 1
						if 'error' in rdata:
							self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['error'], "Table Data", 4)

						if showimages:
							cellcol = 1
							rownum += 1
							self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Screenshot), "Table Heading", 0)

							cellcol += 1

							if 'image_file' in rdata:
								self.xlsx_sections_errors_fill_cell(cellcol, rownum, " ", "Table Data", 4)
								img = openpyxl.drawing.image.Image(rdata['image_file'])
								cellname = ws.cell(row=rownum, column=cellcol).coordinate

								base.debugmsg(5, "img.width:", img.width, "	img.height:", img.height)
								# 								31.75					32.60
								# 								22.23					22.82	==> 70%
								newiw = 850
								ratio = newiw / img.width
								base.debugmsg(5, "ratio:", ratio)
								newih = img.height * ratio
								base.debugmsg(5, "newih:", newih)
								img.width = newiw
								img.height = newih

								newh = newih * 0.76
								# 43.44 cm ==> 32.95 cm	==>	76%
								base.debugmsg(5, "newh:", newh)
								ws.row_dimensions[rownum].height = newh

								ws.add_image(img, cellname)
							else:
								self.xlsx_sections_errors_fill_cell(cellcol, rownum, lbl_NoScreenshot, "Table Data", 4)

		if groupet and not grouprn:
			for errortext in list(grpdata["errortexts"].keys()):
				basekey = grpdata["errortexts"][errortext]["keys"][0]
				base.debugmsg(5, "basekey:", basekey)
				rdata = base.reportdata[id][basekey]

				colspan = 8

				cellcol = 1
				rownum += 1

				self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Error), "Table Heading", 0)
				cellcol += 1
				if 'error' in rdata:
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['error'], "Table Data", colspan)

				cellcol += colspan + 1
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Count), "Table Heading", 0)
				count = len(grpdata["errortexts"][errortext]["keys"])
				cellcol += 1
				self.xlsx_sections_errors_fill_cell(cellcol, rownum, str(count), "Table Data", 0)

				if showimages:
					cellcol = 1
					rownum += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Screenshot), "Table Heading", 0)

					cellcol += 1

					if 'image_file' in rdata:
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, " ", "Table Data", colspan)
						img = openpyxl.drawing.image.Image(rdata['image_file'])
						cellname = ws.cell(row=rownum, column=cellcol).coordinate

						base.debugmsg(5, "img.width:", img.width, "	img.height:", img.height)
						# 									31.75					32.60
						# 									22.23					22.82	==> 70%
						newiw = 850
						ratio = newiw / img.width
						base.debugmsg(5, "ratio:", ratio)
						newih = img.height * ratio
						base.debugmsg(5, "newih:", newih)
						img.width = newiw
						img.height = newih

						newh = newih * 0.76
						# 43.44 cm ==> 32.95 cm	==>	76%
						base.debugmsg(5, "newh:", newh)
						ws.row_dimensions[rownum].height = newh

						ws.add_image(img, cellname)
					else:
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, lbl_NoScreenshot, "Table Data", colspan)

		if not grouprn and not groupet:
			keys = list(base.reportdata[id].keys())
			for key in keys:
				base.debugmsg(5, "key:", key)
				if key != "key":
					rdata = base.reportdata[id][key]

					cellcol = 1
					rownum += 1

					self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Result), "Table Heading", 0)
					cellcol += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['result_name'], "Table Data", 0)

					cellcol += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Test), "Table Heading", 0)
					cellcol += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['test_name'], "Table Data", 0)

					cellcol += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Script), "Table Heading", 0)
					cellcol += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['script'], "Table Data", 0)

					cellcol = 1
					rownum += 1
					self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Error), "Table Heading", 0)
					cellcol += 1
					if 'error' in rdata:
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, rdata['error'], "Table Data", 4)

					if showimages:
						cellcol = 1
						rownum += 1
						self.xlsx_sections_errors_fill_cell(cellcol, rownum, "{}:".format(lbl_Screenshot), "Table Heading", 0)

						cellcol += 1

						if 'image_file' in rdata:
							self.xlsx_sections_errors_fill_cell(cellcol, rownum, " ", "Table Data", 4)
							img = openpyxl.drawing.image.Image(rdata['image_file'])
							cellname = ws.cell(row=rownum, column=cellcol).coordinate

							base.debugmsg(5, "img.width:", img.width, "	img.height:", img.height)
							# 								31.75					32.60
							# 								22.23					22.82	==> 70%
							newiw = 850
							ratio = newiw / img.width
							base.debugmsg(5, "ratio:", ratio)
							newih = img.height * ratio
							base.debugmsg(5, "newih:", newih)
							img.width = newiw
							img.height = newih

							newh = newih * 0.76
							# 43.44 cm ==> 32.95 cm	==>	76%
							base.debugmsg(5, "newh:", newh)
							ws.row_dimensions[rownum].height = newh

							ws.add_image(img, cellname)
						else:
							self.xlsx_sections_errors_fill_cell(cellcol, rownum, lbl_NoScreenshot, "Table Data", 4)

		rownum += 1
		rownum += 1
		self.xlsx_select_cell(1, rownum)

	def xlsx_sections_errors_fill_cell(self, cellcol, rownum, val, style, span):
		base.debugmsg(5, "cellcol:", cellcol, "	rownum:", rownum, "	val:", val, "	style:", style, "	span:", span)
		wb = self.cg_data["xlsx"]["Workbook"]
		ws = wb.active

		self.xlsx_select_cell(cellcol, rownum)

		base.debugmsg(5, "setting Cell value")
		val = base.illegal_xml_chars_re.sub('', val)
		base.debugmsg(8, "val:", val)
		cell = ws.cell(column=cellcol, row=rownum, value=val)
		base.debugmsg(5, "splitting val to lines")
		lines = str(val).splitlines()

		base.debugmsg(5, "lines:", lines)

		if span > 0:
			base.debugmsg(5, "span:", span)
			ws.merge_cells(start_row=rownum, start_column=cellcol, end_row=rownum, end_column=cellcol + span)
		else:
			base.debugmsg(5, "span:", span)
			currw = ws.column_dimensions[cell.column_letter].width
			valw = 0
			for line in lines:
				valw = max(valw, len(line))
			base.debugmsg(5, "currw:", currw, "	valw:", valw)
			neww = max(currw, valw)
			base.debugmsg(5, "neww:", neww)
			ws.column_dimensions[cell.column_letter].width = neww

		if len(lines) > 1:
			base.debugmsg(5, "len(lines):", len(lines))
			# currh = ws.row_dimensions[rownum].height
			# https://stackoverflow.com/questions/32855656/column-and-row-dimensions-in-openpyxl-are-always-none
			currh = 13
			base.debugmsg(5, "currh:", currh)
			newh = currh * len(lines)
			base.debugmsg(5, "newh:", newh)
			ws.row_dimensions[rownum].height = newh

		base.debugmsg(5, "style:", style)
		cell.style = style


class RFSwarm():

	running = True

	def __init__(self):
		while base.running:
			# time.sleep(300)
			time.sleep(1)


base = ReporterBase()

core = ReporterCore()

try:
	core.mainloop()
except KeyboardInterrupt:
	core.on_closing()
except Exception as e:
	base.debugmsg(1, "core.Exception:", e)
	core.on_closing()


# r = RFSwarm_Reporter()
