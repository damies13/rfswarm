import os
import re
import subprocess

import docx
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from robot.api import logger as LOGGER


def read_paragraphs_docx_file(docx_path: str) -> dict:
    """Read paragraphs docx file and return dictionary."""
    doc = Document(docx_path)

    doc_sections: dict = {}
    current_heading = None

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            current_heading = paragraph.text.strip()
            doc_sections[current_heading] = []

        elif paragraph.text.strip():
            if current_heading:
                doc_sections[current_heading].append(paragraph.text.strip())
            else:
                doc_sections.setdefault('Cover', []).append(paragraph.text.strip())

    return doc_sections


def update_docx_with_libreoffice(docx_path: str):
    """Not working currently."""
    command = [
        'libreoffice', '--headless', '--convert-to', 'docx',
        '--outdir', os.path.dirname(docx_path), docx_path
    ]
    process = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if process.returncode != 0:
        raise Exception(f"LibreOffice failed: {process.stderr.decode('utf-8')}")
    else:
        print(f"Document fields updated successfully: {docx_path}")


def update_table_of_contents(docx_path: str):
    """Set updateFields in DOCX file to true."""
    doc = Document(docx_path)

    settings_element = doc.settings.element
    update_fields_element = OxmlElement('w:updateFields')
    update_fields_element.set(qn('w:val'), 'true')

    settings_element.append(update_fields_element)

    doc.save(docx_path)


def get_default_font_name_from_document(docx_path: str) -> str:
    """Returns the name of the font that is used by default in the document."""
    doc = Document(docx_path)

    default_style = doc.styles['Normal']

    if default_style.font.name:
        return default_style.font.name
    else:
        LOGGER.warn("Default font is not set.")
        return 'None'


def extract_docx_images_under_heading(heading: str, docx_path: str, output_folder: str, debug=False) -> list:
    """
    Extract images from a DOCX file, including images inside tables,
    and associate them to the provided heading.
    """
    doc = Document(docx_path)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    image_counter = 1
    current_heading_name = None
    image_list = []

    def extract_images_from_run(run, heading, output_folder):
        """Extract images from a run and save them to the given output folder."""
        nonlocal image_counter

        if run.element.xpath('.//w:drawing'):
            for rel in run.element.xpath('.//a:blip/@r:embed'):
                part = doc.part.related_parts[rel]

                image_name = f"{heading.replace(' ', '_')}_{image_counter}_image.png"
                image_path = os.path.join(output_folder, image_name)

                with open(image_path, 'wb') as img_file:
                    img_file.write(part.blob)

                print(f"DOCX RFSwarm report image saved as: {image_name}") if debug else 0
                image_list.append(image_name)
                image_counter += 1

    for inner_obj in doc.iter_inner_content():
        image_counter = 1

        if isinstance(inner_obj, docx.text.paragraph.Paragraph):
            paragraph = inner_obj
            if paragraph.style.name.startswith("Heading"):
                heading_name = paragraph.text.strip()
                heading_name_splitted = heading_name.split(" ", maxsplit=1)
                current_heading_name = heading_name_splitted[1]

            if current_heading_name == heading:
                for run in paragraph.runs:
                    extract_images_from_run(run, current_heading_name, output_folder)

        elif isinstance(inner_obj, docx.table.Table):
            if current_heading_name == heading:
                table = inner_obj

                def iter_unique_row_cells(row):
                    """Generate cells in given row skipping empty grid cells."""
                    last_cell_tc = None
                    for cell in row.cells:
                        this_cell_tc = cell._tc
                        if this_cell_tc is last_cell_tc:
                            continue
                        last_cell_tc = this_cell_tc
                        yield cell

                for row in table.rows:
                    for cell in iter_unique_row_cells(row):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                extract_images_from_run(run, current_heading_name, output_folder)

    return image_list


def extract_images_from_docx(docx_path: str, output_folder: str) -> list:
    """
    Extract all images from DOCX one by one.
    Returns list of the saved images.
    """
    doc = Document(docx_path)
    saved_images = []

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_name = os.path.basename(rel.target_ref)
            img_data = rel.target_part.blob

            with open(os.path.join(output_folder, img_name), "wb") as img_file:
                img_file.write(img_data)
                print(f"DOCX RFSwarm report image saved as: {img_name}")
                saved_images.append(img_name)

    return saved_images


def extract_table_from_docx(docx_table_obj: docx.table.Table) -> list:
    """
    Extract table from docx.table.Table object as nested list.
    """

    def iter_unique_row_cells(row):
        """Generate cells in given row skipping empty grid cells."""
        last_cell_tc = None
        for cell in row.cells:
            this_cell_tc = cell._tc
            if this_cell_tc is last_cell_tc:
                continue
            last_cell_tc = this_cell_tc
            yield cell

    table_data = []
    for row in docx_table_obj.rows:
        row_data = []
        for cell in iter_unique_row_cells(row):
            row_data.append(cell.text)
        table_data.append(row_data)

    return table_data


def sort_docx_images(docx_images: list) -> list:
    """Sort given list using natural sorting"""
    def natural_sort_key(filename):
        return [int(text) if text.isdigit() else text for text in re.split(r'(\d+)', filename)]

    return sorted(docx_images, key=natural_sort_key)


def read_docx_file(docx_path: str, debug=False) -> dict:
    """
    Read the docx file in terms of text and table content.
    Saved data is saved as:
    dict = {"Heading 1" = {"text" = ["line 1", "line 2", ...], "table" = [[row 1], [row 2], ...] }, ...}
    """
    doc = Document(docx_path)
    doc_data: dict = {}

    current_heading_name = None

    for inner_obj in doc.iter_inner_content():
        if isinstance(inner_obj, docx.text.paragraph.Paragraph):
            paragraph = inner_obj
            if paragraph.style.name.startswith('Heading'):
                heading_name = paragraph.text.strip()
                heading_name_splitted = heading_name.split(" ", maxsplit=1)
                current_heading_name = heading_name_splitted[1]
                doc_data.setdefault(current_heading_name, {"text": [], "table": []})

            elif paragraph.text.strip():
                if current_heading_name:
                    doc_data[current_heading_name]["text"].append(paragraph.text.strip())
                else:
                    doc_data.setdefault("Cover", {"text": []})
                    doc_data["Cover"]["text"].append(paragraph.text.strip())

        elif isinstance(inner_obj, docx.table.Table):
            table = inner_obj
            table_list = extract_table_from_docx(table)
            doc_data[current_heading_name]["table"] = table_list

    if debug:
        for k in doc_data:
            print(" -", k, ":")
            print(doc_data[k], end="\n\n\n")

    return doc_data
